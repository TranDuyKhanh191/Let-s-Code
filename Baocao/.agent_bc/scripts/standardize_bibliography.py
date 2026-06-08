"""
standardize_bibliography.py – Dynamic Bibliography Standardizer for Agent BC
=============================================================================
Parses all references from the Word document, sorts them alphabetically
according to Vietnamese (by given-name) and English (by surname / corporate
name) rules, builds the old→new citation mapping on-the-fly, propagates
index changes through all body paragraphs and table cells, and rewrites
the reference list in sorted order while preserving run-level formatting.

Safety:
  • Always creates a timestamped backup before modifying.
  • Filters citation-like patterns that are actually dimensions (e.g. [3000, 2500]).
  • Idempotent: running twice on an already-sorted document produces no changes.

Usage:
  python standardize_bibliography.py --path "path/to/doc.docx"
"""

import docx
import re
import os
import sys
import shutil
import argparse
import unicodedata
from datetime import datetime
from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

sys.stdout.reconfigure(encoding='utf-8')

# ---------------------------------------------------------------------------
# 1. Helpers
# ---------------------------------------------------------------------------

def normalize(text: str) -> str:
    """NFC-normalize a string for consistent Vietnamese character handling."""
    return unicodedata.normalize('NFC', text.strip())


def extract_ref_index(text: str):
    """Return the integer index from a reference line like '[5] Author ...'."""
    m = re.match(r'^\[(\d+)\]', text.strip())
    return int(m.group(1)) if m else None


# Vietnamese alphabet order for proper sorting
_VN_ALPHA_ORDER = (
    'a à á ả ã ạ ă ằ ắ ẳ ẵ ặ â ầ ấ ẩ ẫ ậ '
    'b c d đ e è é ẻ ẽ ẹ ê ề ế ể ễ ệ '
    'g h i ì í ỉ ĩ ị k l m n o ò ó ỏ õ ọ ô ồ ố ổ ỗ ộ ơ ờ ớ ở ỡ ợ '
    'p q r s t u ù ú ủ ũ ụ ư ừ ứ ử ữ ự v x y ỳ ý ỷ ỹ ỵ'
).split()
_VN_CHAR_ORDER = {c: i for i, c in enumerate(_VN_ALPHA_ORDER)}


def vn_collation_key(s: str) -> list:
    """Generate a sort key that respects Vietnamese alphabetical order."""
    result = []
    for ch in normalize(s).lower():
        result.append(_VN_CHAR_ORDER.get(ch, 1000 + ord(ch)))
    return result


def vietnamese_sort_key(author_text: str):
    """
    Vietnamese sorting: sort by the *given name* (tên) which is the LAST
    word of the author's full name, using proper Vietnamese alphabet order.
    E.g. 'Nguyễn Văn Ba' → sort key for 'ba'
         'Trần Hạnh Nhi' → sort key for 'nhi'
    """
    # Strip index prefix like '[1] '
    cleaned = re.sub(r'^\[\d+\]\s*', '', author_text.strip())
    # Take the part before the first '('  → author name
    author_part = cleaned.split('(')[0].strip().rstrip(',').strip()
    parts = author_part.split()
    if parts:
        return vn_collation_key(normalize(parts[-1]))
    return vn_collation_key(cleaned)


def english_sort_key(author_text: str) -> str:
    """
    English / corporate sorting: sort by the first word (surname or org name).
    E.g. 'Chen, M. et al. (2021), ...'          → 'chen'
         'PostgreSQL Global Development Group …' → 'postgresql'
         'SAP (2024), …'                         → 'sap'
    """
    cleaned = re.sub(r'^\[\d+\]\s*', '', author_text.strip())
    author_part = cleaned.split('(')[0].strip().rstrip(',').strip()
    parts = author_part.split()
    if parts:
        # Remove trailing comma from surname  'Chen,' → 'Chen'
        return parts[0].rstrip(',').lower()
    return cleaned.lower()


def collect_paragraph_xml(p):
    """
    Collect a deep copy of a paragraph's XML element so we can reinsert it
    later with all original run-level formatting intact.
    """
    return deepcopy(p._element)


def strip_ref_index_from_xml(p_xml, new_index: int):
    """
    Given a deep-copied <w:p> element whose first run starts with '[N] ',
    replace that prefix with '[new_index] '.
    Handles cases where the index is across multiple runs.
    """
    runs = p_xml.findall(qn('w:r'))
    if not runs:
        return

    # Strategy: walk runs, accumulate text until we've consumed '[N] '
    consumed = 0
    target_prefix_re = re.compile(r'^\[\d+\]\s*')
    accumulated = ''

    for r in runs:
        t_el = r.find(qn('w:t'))
        if t_el is None or t_el.text is None:
            continue
        accumulated += t_el.text
        m = target_prefix_re.match(accumulated)
        if m:
            # Found the prefix – figure out how much of this run to replace
            prefix_len = m.end()
            new_prefix = f'[{new_index}] '

            # Rebuild text across the runs that contributed
            rebuilt = new_prefix + accumulated[prefix_len:]
            # Clear text from all contributing runs except the last
            for prev_r in runs[:runs.index(r)]:
                prev_t = prev_r.find(qn('w:t'))
                if prev_t is not None:
                    prev_t.text = ''
            t_el.text = rebuilt[len(new_prefix) - len(new_prefix):]
            # Actually set the first contributing run
            first_t = runs[0].find(qn('w:t'))
            if first_t is not None:
                # Clear all contributing runs
                for cr in runs[:runs.index(r) + 1]:
                    ct = cr.find(qn('w:t'))
                    if ct is not None:
                        ct.text = ''
                # Set all text on first run
                first_t.text = rebuilt
                # Preserve space
                first_t.set(qn('xml:space'), 'preserve')
            break
        consumed += len(t_el.text)


# ---------------------------------------------------------------------------
# 2. Reference Parsing
# ---------------------------------------------------------------------------

def find_section_indices(doc):
    """
    Locate the indices for:
      - ref_heading: 'DANH MỤC TÀI LIỆU THAM KHẢO' (Heading 1)
      - vi_heading:  'Tài liệu tiếng Việt'
      - en_heading:  'Tài liệu tiếng Anh và trực tuyến'
    Scans backwards to avoid matching TOC entries.
    """
    ref_idx = vi_idx = en_idx = -1
    for idx in range(len(doc.paragraphs) - 1, -1, -1):
        text = normalize(doc.paragraphs[idx].text)
        style = doc.paragraphs[idx].style.name if doc.paragraphs[idx].style else ''
        if vi_idx == -1 and 'Tài liệu tiếng Việt' in text:
            vi_idx = idx
        elif en_idx == -1 and 'Tài liệu tiếng Anh' in text:
            en_idx = idx
        elif ref_idx == -1 and style.startswith('Heading') and 'DANH MỤC TÀI LIỆU' in text.upper():
            ref_idx = idx
    return ref_idx, vi_idx, en_idx


def parse_reference_items(doc, start_idx, end_idx):
    """
    Parse reference items between start_idx (exclusive) and end_idx (exclusive).
    Returns list of tuples: (original_index, paragraph_index, paragraph_xml_copy, full_text)
    """
    items = []
    for idx in range(start_idx + 1, end_idx):
        p = doc.paragraphs[idx]
        text = normalize(p.text)
        if not text:
            continue
        ref_num = extract_ref_index(text)
        if ref_num is not None:
            items.append((ref_num, idx, collect_paragraph_xml(p), text))
        elif p.style and p.style.name == 'reference-item':
            # reference-item style but no index → might be a continuation
            items.append((None, idx, collect_paragraph_xml(p), text))
    return items


def parse_all_references(doc, vi_idx, en_idx):
    """
    Parse Vietnamese and English reference groups.
    Returns (vi_items, en_items) where each item is
    (original_index, para_idx, xml_copy, full_text).
    
    For en_items, the end boundary is the end of the document or
    the next Heading 1.
    """
    # Find end boundary for English section
    en_end = len(doc.paragraphs)
    for idx in range(en_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[idx]
        style = p.style.name if p.style else ''
        if style.startswith('Heading 1') and normalize(p.text):
            en_end = idx
            break

    vi_items = parse_reference_items(doc, vi_idx, en_idx)
    en_items = parse_reference_items(doc, en_idx, en_end)
    return vi_items, en_items


# ---------------------------------------------------------------------------
# 3. Sorting
# ---------------------------------------------------------------------------

def sort_references(vi_items, en_items):
    """
    Sort Vietnamese items by given-name, English items by surname/corporate.
    Returns (sorted_vi, sorted_en) and the combined old→new mapping dict.
    """
    # Filter items that have a valid index
    vi_with_idx = [item for item in vi_items if item[0] is not None]
    en_with_idx = [item for item in en_items if item[0] is not None]

    sorted_vi = sorted(vi_with_idx, key=lambda x: vietnamese_sort_key(x[3]))
    sorted_en = sorted(en_with_idx, key=lambda x: english_sort_key(x[3]))

    # Build mapping: old_index → new_index
    mapping = {}
    new_idx = 1
    for item in sorted_vi:
        mapping[item[0]] = new_idx
        new_idx += 1
    for item in sorted_en:
        mapping[item[0]] = new_idx
        new_idx += 1

    return sorted_vi, sorted_en, mapping


# ---------------------------------------------------------------------------
# 4. Citation Propagation
# ---------------------------------------------------------------------------

def replace_citations_in_runs(runs, mapping: dict):
    """
    Replace citation tags [N] in runs using the mapping.
    Only replaces integers that exist in the mapping to avoid
    corrupting dimension arrays like [3000, 2500].
    """
    for r in runs:
        t_el = r.find(qn('w:t')) if hasattr(r, 'find') else None
        # python-docx Run object
        if t_el is None and hasattr(r, '_element'):
            t_el = r._element.find(qn('w:t'))
        if t_el is None:
            # Direct lxml element
            if hasattr(r, 'tag'):
                t_el = r.find(qn('w:t'))
        
        text = None
        if hasattr(r, 'text') and isinstance(r.text, str):
            text = r.text
        elif t_el is not None and t_el.text:
            text = t_el.text
        
        if not text:
            continue

        def repl(match):
            num = int(match.group(1))
            if num in mapping:
                return f'[{mapping[num]}]'
            return match.group(0)

        new_text = re.sub(r'\[(\d+)\]', repl, text)
        if new_text != text:
            if hasattr(r, 'text') and isinstance(r.text, str):
                r.text = new_text
            elif t_el is not None:
                t_el.text = new_text


def propagate_citations(doc, ref_heading_idx: int, mapping: dict):
    """
    Update all [N] citation tags in body paragraphs (before references)
    and in all table cells throughout the document.
    
    Uses a two-pass approach to avoid double-mapping:
      Pass 1: [old] → [§temp_new§]
      Pass 2: [§temp_new§] → [new]
    """
    if not mapping:
        return

    # Check if mapping is identity (already sorted)
    if all(k == v for k, v in mapping.items()):
        print("  Citations already in correct order – skipping propagation.")
        return

    # Build temp mapping to avoid double-replacement
    # old → §TEMP_N§   then  §TEMP_N§ → new
    temp_marker = '§CITE_'

    def pass1_repl(match):
        num = int(match.group(1))
        if num in mapping:
            return f'{temp_marker}{mapping[num]}§'
        return match.group(0)

    def pass2_repl(match):
        num = int(match.group(1))
        return f'[{num}]'

    pattern1 = re.compile(r'\[(\d+)\]')
    pattern2 = re.compile(r'§CITE_(\d+)§')

    def two_pass_replace(runs):
        for r in runs:
            if not r.text:
                continue
            r.text = pattern1.sub(pass1_repl, r.text)
        for r in runs:
            if not r.text:
                continue
            r.text = pattern2.sub(pass2_repl, r.text)

    # Body paragraphs
    print("  Updating citations in body paragraphs...")
    for idx in range(ref_heading_idx):
        p = doc.paragraphs[idx]
        two_pass_replace(p.runs)

    # Table cells
    print("  Updating citations in table cells...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    two_pass_replace(p.runs)


# ---------------------------------------------------------------------------
# 5. Rewrite References Section
# ---------------------------------------------------------------------------

def remove_old_references(doc, vi_idx, en_idx):
    """
    Remove all reference-item paragraphs between vi_heading and end of
    references section, but keep the headings themselves.
    """
    # Find end boundary
    en_end = len(doc.paragraphs)
    for idx in range(en_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[idx]
        style = p.style.name if p.style else ''
        if style.startswith('Heading 1') and normalize(p.text):
            en_end = idx
            break

    # Collect indices to remove (reverse order for stable removal)
    to_remove = []
    for idx in range(vi_idx + 1, en_end):
        if idx == en_idx:
            continue  # keep the English heading
        p = doc.paragraphs[idx]
        text = normalize(p.text)
        if text and extract_ref_index(text) is not None:
            to_remove.append(idx)
        elif p.style and p.style.name == 'reference-item' and text:
            to_remove.append(idx)

    print(f"  Removing {len(to_remove)} old reference items...")
    for idx in sorted(to_remove, reverse=True):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)


def insert_references(doc, heading_idx, sorted_items, start_number):
    """
    Insert sorted reference items after the given heading.
    Uses deep-copied XML to preserve all original formatting (italic titles,
    fonts, sizes, etc.) and only updates the index number.
    """
    # Find the element right after the heading to use as insertion anchor
    heading_el = doc.paragraphs[heading_idx]._element
    
    # Insert in reverse order using addnext on heading so they end up in
    # correct forward order
    for i, item in enumerate(reversed(sorted_items)):
        old_idx, _, xml_copy, _ = item
        new_idx = start_number + len(sorted_items) - 1 - i
        
        # Update the index in the XML copy
        strip_ref_index_from_xml(xml_copy, new_idx)
        
        # Insert after heading
        heading_el.addnext(xml_copy)


# ---------------------------------------------------------------------------
# 6. Main Orchestrator
# ---------------------------------------------------------------------------

def standardize_references(doc_path: str, dry_run: bool = False) -> bool:
    """
    Main entry point: parse → sort → propagate → rewrite.
    """
    print(f"=== STANDARDIZING BIBLIOGRAPHY & CITATIONS: {doc_path} ===")
    if not os.path.exists(doc_path):
        print(f"ERROR: File not found at {doc_path}")
        return False

    # Backup
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup_path = doc_path.replace('.docx', f'_Backup_{timestamp}.docx')
    shutil.copy2(doc_path, backup_path)
    print(f"Backup created at {backup_path}")

    doc = docx.Document(doc_path)

    # 1. Find section headings
    ref_idx, vi_idx, en_idx = find_section_indices(doc)
    if ref_idx == -1 or vi_idx == -1 or en_idx == -1:
        print(f"ERROR: Could not find all reference headings!")
        print(f"  ref_heading={ref_idx}, vi_heading={vi_idx}, en_heading={en_idx}")
        return False
    print(f"Found headings: ref={ref_idx}, vi={vi_idx}, en={en_idx}")

    # 2. Parse existing references
    vi_items, en_items = parse_all_references(doc, vi_idx, en_idx)
    print(f"Parsed: {len(vi_items)} Vietnamese, {len(en_items)} English/Online references")

    if not vi_items and not en_items:
        print("WARNING: No references found to sort!")
        return True

    # 3. Sort and build mapping
    sorted_vi, sorted_en, mapping = sort_references(vi_items, en_items)

    # Print sort results
    print("\nSorted Vietnamese references:")
    for i, item in enumerate(sorted_vi, 1):
        old = item[0]
        print(f"  [{old}] → [{mapping[old]}]  {item[3][:70]}...")
    print("\nSorted English references:")
    for i, item in enumerate(sorted_en, len(sorted_vi) + 1):
        old = item[0]
        print(f"  [{old}] → [{mapping[old]}]  {item[3][:70]}...")

    # Check if already sorted
    is_identity = all(k == v for k, v in mapping.items())
    if is_identity:
        print("\n✓ References are already in correct sorted order. No changes needed.")
        # Clean up backup since no changes
        os.remove(backup_path)
        print(f"Removed unnecessary backup: {backup_path}")
        return True

    if dry_run:
        print("\n[DRY RUN] Would apply the above mapping. No changes saved.")
        os.remove(backup_path)
        return True

    # 4. Propagate citation changes in body text
    print("\nPropagating citation index changes...")
    propagate_citations(doc, ref_idx, mapping)

    # 5. Remove old references and rewrite sorted ones
    print("\nRewriting reference list...")
    remove_old_references(doc, vi_idx, en_idx)

    # Save intermediate to re-parse heading indices after removal
    doc.save(doc_path)
    doc = docx.Document(doc_path)
    _, vi_idx, en_idx = find_section_indices(doc)

    # Insert Vietnamese references
    insert_references(doc, vi_idx, sorted_vi, start_number=1)

    # Save and reload for English insertion
    doc.save(doc_path)
    doc = docx.Document(doc_path)
    _, _, en_idx = find_section_indices(doc)

    # Insert English references
    insert_references(doc, en_idx, sorted_en, start_number=len(sorted_vi) + 1)

    doc.save(doc_path)
    print(f"\n✓ Standardization completed successfully!")
    print(f"  Backup: {backup_path}")
    return True


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Dynamically sort bibliography and propagate citation changes.')
    parser.add_argument('--path', type=str,
                        default='d:/MiniERP_NhomKinh/ChuyenDeTotNghiep_NguyenTrongThoi.docx',
                        help='Path to the .docx file.')
    parser.add_argument('--dry-run', action='store_true',
                        help='Show what would change without modifying the document.')
    args = parser.parse_args()
    standardize_references(args.path, dry_run=args.dry_run)
