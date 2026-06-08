"""
audit_doc_integrity.py – Comprehensive Word Document Integrity Auditor
======================================================================
Performs a full audit of a Word (.docx) document to catch formatting
issues, citation inconsistencies, draft placeholders, and configuration
problems before submission.

Audit Checks:
  [1] Duplicate reference index detection
  [2] Orphan references (listed but never cited in body)
  [3] Phantom citations (cited in body but not listed in references)
  [4] Citation coverage completeness
  [5] Draft placeholder and risk keyword scan
  [6] Figure/table caption formatting (non-bold, correct placement)
  [7] Reference alphabetical sorting verification
  [8] TOC auto-update configuration (w:updateFields)

Usage:
  python audit_doc_integrity.py --path "path/to/doc.docx"
  python audit_doc_integrity.py --path "path/to/doc.docx" --strict
"""

import docx
import sys
import re
import os
import unicodedata
import argparse
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def normalize(text: str) -> str:
    return unicodedata.normalize('NFC', text.strip())


# Vietnamese alphabet order for proper sorting
_VN_ALPHA_ORDER = (
    'a à á ả ã ạ ă ằ ắ ẳ ẵ ặ â ầ ấ ẩ ẫ ậ '
    'b c d đ e è é ẻ ẽ ẹ ê ề ế ể ễ ệ '
    'g h i ì í ỉ ĩ ị k l m n o ò ó ỏ õ ọ ô ồ ố ổ ỗ ộ ơ ờ ớ ở ỡ ợ '
    'p q r s t u ù ú ủ ũ ụ ư ừ ứ ử ữ ự v x y ỳ ý ỷ ỹ ỵ'
).split()
_VN_CHAR_ORDER = {c: i for i, c in enumerate(_VN_ALPHA_ORDER)}


def vn_collation_key(s: str) -> list:
    """Generate a sort key that respects Vietnamese alphabetical order."""
    result = []
    for ch in normalize(s).lower():
        result.append(_VN_CHAR_ORDER.get(ch, 1000 + ord(ch)))
    return result


def vietnamese_sort_key(text: str):
    """Sort key for Vietnamese authors: by given name (last word), using proper VN collation."""
    cleaned = re.sub(r'^\[\d+\]\s*', '', text.strip())
    author_part = cleaned.split('(')[0].strip().rstrip(',').strip()
    parts = author_part.split()
    name = normalize(parts[-1]).lower() if parts else cleaned.lower()
    return vn_collation_key(name)


def english_sort_key(text: str) -> str:
    """Sort key for English/corporate authors: by first word (surname)."""
    cleaned = re.sub(r'^\[\d+\]\s*', '', text.strip())
    author_part = cleaned.split('(')[0].strip().rstrip(',').strip()
    parts = author_part.split()
    return parts[0].rstrip(',').lower() if parts else cleaned.lower()


# ---------------------------------------------------------------------------
# Main Audit
# ---------------------------------------------------------------------------

def audit_document(doc_path: str, strict: bool = False) -> bool:
    print(f"{'='*60}")
    print(f"  DOCUMENT INTEGRITY AUDIT")
    print(f"  {doc_path}")
    print(f"{'='*60}\n")

    if not os.path.exists(doc_path):
        print(f"FATAL: File not found at {doc_path}")
        return False

    doc = docx.Document(doc_path)
    issues = []  # Collect all issues
    warnings = []

    # -----------------------------------------------------------------------
    # 0. Locate References Section
    # -----------------------------------------------------------------------
    ref_idx = vi_idx = en_idx = -1
    for idx in range(len(doc.paragraphs) - 1, -1, -1):
        text = normalize(doc.paragraphs[idx].text)
        style = doc.paragraphs[idx].style.name if doc.paragraphs[idx].style else ''
        if vi_idx == -1 and 'Tài liệu tiếng Việt' in text:
            vi_idx = idx
        elif en_idx == -1 and 'Tài liệu tiếng Anh' in text:
            en_idx = idx
        elif ref_idx == -1 and style.startswith('Heading') and 'DANH MỤC TÀI LIỆU' in text.upper():
            ref_idx = idx

    if ref_idx == -1:
        # Fallback: search for TÀI LIỆU THAM KHẢO in any text
        for idx in range(len(doc.paragraphs) - 1, -1, -1):
            text = normalize(doc.paragraphs[idx].text)
            if 'TÀI LIỆU THAM KHẢO' in text.upper():
                ref_idx = idx
                print(f"  ℹ Found fallback references heading at P[{ref_idx}]")
                break

    if ref_idx == -1:
        print("[0] FATAL: References section heading not found!")
        return False
    else:
        print(f"[0] References section found at P[{ref_idx}]")
        if vi_idx != -1:
            print(f"    Vietnamese sub-heading at P[{vi_idx}]")
        if en_idx != -1:
            print(f"    English sub-heading at P[{en_idx}]")

    # -----------------------------------------------------------------------
    # 1. Parse reference items and check duplicates
    # -----------------------------------------------------------------------
    print(f"\n[1] Checking for duplicate reference indices...")
    ref_items = []  # (index, text, paragraph_idx)
    for idx in range(ref_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[idx]
        text = normalize(p.text)
        if not text:
            continue
        # Stop if we hit another Heading 1 that isn't a sub-heading
        style = p.style.name if p.style else ''
        if style.startswith('Heading 1') and text and idx != vi_idx and idx != en_idx:
            # Check if this is part of references or a new chapter
            if 'TÀI LIỆU' not in text.upper():
                break
        m = re.match(r'^\[(\d+)\]', text)
        if m:
            ref_items.append((int(m.group(1)), text, idx))

    ref_numbers = [r[0] for r in ref_items]
    print(f"    Parsed {len(ref_numbers)} reference items.")

    duplicates = set(x for x in ref_numbers if ref_numbers.count(x) > 1)
    if duplicates:
        msg = f"Duplicate reference indices: {sorted(duplicates)}"
        issues.append(msg)
        print(f"    ✗ {msg}")
    else:
        print(f"    ✓ No duplicate reference numbers.")

    # Check for sequential numbering
    expected = list(range(1, len(ref_numbers) + 1))
    if ref_numbers != expected:
        if sorted(ref_numbers) == expected:
            warnings.append(f"References are not in sequential order: {ref_numbers}")
            print(f"    ⚠ References not in sequential 1..{len(ref_numbers)} order")
        else:
            gaps = set(expected) - set(ref_numbers)
            extra = set(ref_numbers) - set(expected)
            if gaps:
                warnings.append(f"Missing reference numbers: {sorted(gaps)}")
                print(f"    ⚠ Missing numbers: {sorted(gaps)}")
            if extra:
                warnings.append(f"Unexpected reference numbers: {sorted(extra)}")
                print(f"    ⚠ Unexpected numbers: {sorted(extra)}")
    else:
        print(f"    ✓ Sequential numbering [1]-[{len(ref_numbers)}] confirmed.")

    # -----------------------------------------------------------------------
    # 2. Citation coverage analysis
    # -----------------------------------------------------------------------
    print(f"\n[2] Analyzing citation coverage...")

    # Collect all citations from body text (before references)
    body_citations = set()
    for idx in range(ref_idx):
        text = doc.paragraphs[idx].text
        for m in re.finditer(r'\[(\d+)\]', text):
            num = int(m.group(1))
            if num <= 100:  # Filter out dimension-like numbers
                body_citations.add(num)

    # Also check table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for m in re.finditer(r'\[(\d+)\]', p.text):
                        num = int(m.group(1))
                        if num <= 100:
                            body_citations.add(num)

    ref_set = set(ref_numbers)
    total_cite_count = 0
    for idx in range(ref_idx):
        total_cite_count += len(re.findall(r'\[\d+\]', doc.paragraphs[idx].text))

    print(f"    Total citation occurrences in body: {total_cite_count}")
    print(f"    Unique citations in body: {sorted(body_citations)}")
    print(f"    Listed references: {sorted(ref_set)}")

    # Orphan references (listed but not cited)
    orphans = ref_set - body_citations
    if orphans:
        msg = f"Orphan references (listed but never cited): {sorted(orphans)}"
        if strict:
            issues.append(msg)
            print(f"    ✗ {msg}")
        else:
            warnings.append(msg)
            print(f"    ⚠ {msg}")
    else:
        print(f"    ✓ All listed references are cited in body text.")

    # Phantom citations (cited but not listed)
    phantoms = body_citations - ref_set
    if phantoms:
        msg = f"Phantom citations (cited but not in reference list): {sorted(phantoms)}"
        issues.append(msg)
        print(f"    ✗ {msg}")
    else:
        print(f"    ✓ All body citations have matching reference entries.")

    # Full coverage
    if body_citations == ref_set:
        print(f"    ✓ Perfect citation coverage: {len(ref_set)} references, all cited.")

    # -----------------------------------------------------------------------
    # 3. Draft placeholders and risk keywords
    # -----------------------------------------------------------------------
    print(f"\n[3] Scanning for draft placeholders and risk keywords...")
    risk_keywords = [
        "[placeholder]", "chưa có", "copy từ", "wikipedia", "theo nguồn",
        "chưa rõ", "ảnh ở đây", "placeholder_code",
        "chú thích hình", "cần thêm", "tự viết", "xem thêm ở",
        "TODO", "FIXME", "HACK", "draft",
        "lorem ipsum", "sample text", "example text",
        "insert here", "add later", "TBD", "hình minh họa"
    ]
    # False positive patterns to skip
    false_positive_patterns = [
        "chưa có tài khoản",
        "chưa có đơn hàng",
        "chưa có dữ liệu",
        "chưa có phân công",
        "-xxx",  # Document codes like DH-xxx, PC-xxx
    ]

    found_risks = 0
    search_limit = ref_idx
    for idx in range(search_limit):
        p = doc.paragraphs[idx]
        text = normalize(p.text)
        text_lower = text.lower()
        for kw in risk_keywords:
            if kw.lower() in text_lower:
                # Check false positives
                is_fp = any(fp.lower() in text_lower for fp in false_positive_patterns)
                if is_fp:
                    continue
                preview = text[:100] + ('...' if len(text) > 100 else '')
                warnings.append(f"P[{idx}] keyword '{kw}': {preview}")
                print(f"    ⚠ P[{idx}] (keyword: '{kw}'): '{preview}'")
                found_risks += 1

    if found_risks == 0:
        print(f"    ✓ No draft keywords or placeholders found.")
    else:
        print(f"    Found {found_risks} potential risk areas.")

    # -----------------------------------------------------------------------
    # 4. Figure and table caption formatting
    # -----------------------------------------------------------------------
    print(f"\n[4] Auditing figure and table captions...")
    captions = []
    bold_captions = []
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith('Hình ') or text.startswith('Bảng '):
            captions.append((idx, text))
            for r in p.runs:
                if r.font.bold or (r.style and r.style.font and r.style.font.bold):
                    bold_captions.append((idx, text[:80]))
                    break

    print(f"    Found {len(captions)} captioned figures/tables.")
    if bold_captions:
        msg = f"Found {len(bold_captions)} captions with bold text (should be normal weight)"
        warnings.append(msg)
        print(f"    ⚠ {msg}:")
        for idx, text in bold_captions[:5]:
            print(f"      P[{idx}]: '{text}'")
        if len(bold_captions) > 5:
            print(f"      ... and {len(bold_captions) - 5} more")
    else:
        print(f"    ✓ All figure/table captions are in normal weight (non-bold).")

    # -----------------------------------------------------------------------
    # 5. Reference alphabetical sorting verification
    # -----------------------------------------------------------------------
    print(f"\n[5] Verifying reference alphabetical sorting...")
    if vi_idx != -1 and en_idx != -1:
        # Parse Vietnamese refs
        vi_refs = [(r[0], r[1]) for r in ref_items if r[2] > vi_idx and r[2] < en_idx]
        # Parse English refs
        en_refs = [(r[0], r[1]) for r in ref_items if r[2] > en_idx]

        # Check Vietnamese sorting
        vi_keys = [vietnamese_sort_key(r[1]) for r in vi_refs]
        vi_sorted = sorted(vi_keys)
        if vi_keys == vi_sorted:
            print(f"    ✓ Vietnamese references ({len(vi_refs)}) are correctly sorted by given name.")
        else:
            msg = f"Vietnamese references are NOT sorted correctly by given name"
            warnings.append(msg)
            print(f"    ⚠ {msg}")
            for i, (actual, expected) in enumerate(zip(vi_keys, vi_sorted)):
                if actual != expected:
                    print(f"      Position {i+1}: got '{actual}', expected '{expected}'")

        # Check English sorting
        en_keys = [english_sort_key(r[1]) for r in en_refs]
        en_sorted = sorted(en_keys)
        if en_keys == en_sorted:
            print(f"    ✓ English references ({len(en_refs)}) are correctly sorted by surname/org.")
        else:
            msg = f"English references are NOT sorted correctly by surname/org name"
            warnings.append(msg)
            print(f"    ⚠ {msg}")
            for i, (actual, expected) in enumerate(zip(en_keys, en_sorted)):
                if actual != expected:
                    print(f"      Position {i+1}: got '{actual}', expected '{expected}'")
    else:
        print(f"    ⚠ Cannot verify sorting: sub-headings not found.")

    # -----------------------------------------------------------------------
    # 6. TOC auto-update setting
    # -----------------------------------------------------------------------
    print(f"\n[6] Auditing Table of Contents settings...")
    try:
        settings_element = doc.settings.element
        update_fields = settings_element.findall(qn('w:updateFields'))
        if update_fields:
            val = update_fields[0].get(qn('w:val'), 'true')
            print(f"    ✓ w:updateFields is set to '{val}'.")
        else:
            msg = "w:updateFields is NOT configured – TOC will not auto-update on open"
            warnings.append(msg)
            print(f"    ⚠ {msg}")
    except Exception as e:
        print(f"    ⚠ Could not read document settings: {e}")

    # -----------------------------------------------------------------------
    # 7. Soft line break detection (preview only)
    # -----------------------------------------------------------------------
    print(f"\n[7] Scanning for justified paragraphs with soft line breaks...")
    soft_break_count = 0
    for idx, p in enumerate(doc.paragraphs):
        if idx >= ref_idx:
            break
        alignment = p.paragraph_format.alignment
        is_just = (alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY)
        if alignment is None and p.style:
            is_just = (p.style.paragraph_format.alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY)
        if is_just and '\n' in p.text:
            soft_break_count += 1
            if soft_break_count <= 3:
                preview = p.text[:60].replace('\n', '\\n')
                print(f"    P[{idx}]: '{preview}...'")

    if soft_break_count > 0:
        msg = f"Found {soft_break_count} justified paragraphs with soft breaks (run fix_spacing_and_breaks.py)"
        warnings.append(msg)
        print(f"    ⚠ {msg}")
    else:
        print(f"    ✓ No soft line break issues detected.")

    # -----------------------------------------------------------------------
    # Summary
    # -----------------------------------------------------------------------
    print(f"\n{'='*60}")
    print(f"  AUDIT SUMMARY")
    print(f"{'='*60}")
    print(f"  Total issues (errors):   {len(issues)}")
    print(f"  Total warnings:          {len(warnings)}")

    if issues:
        print(f"\n  ERRORS (must fix):")
        for i, issue in enumerate(issues, 1):
            print(f"    {i}. {issue}")

    if warnings:
        print(f"\n  WARNINGS (review recommended):")
        for i, w in enumerate(warnings, 1):
            print(f"    {i}. {w}")

    if not issues and not warnings:
        print(f"\n  ✓✓✓ DOCUMENT PASSES ALL CHECKS ✓✓✓")

    status = "PASS" if not issues else "FAIL"
    print(f"\n  Final verdict: {status}")
    print(f"{'='*60}\n")

    return len(issues) == 0


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Audit Word Document integrity.')
    parser.add_argument('--path', type=str,
                        default='d:/MiniERP_NhomKinh/ChuyenDeTotNghiep_NguyenTrongThoi.docx',
                        help='Path to the .docx file.')
    parser.add_argument('--strict', action='store_true',
                        help='Treat orphan references as errors instead of warnings.')
    args = parser.parse_args()
    audit_document(args.path, strict=args.strict)
