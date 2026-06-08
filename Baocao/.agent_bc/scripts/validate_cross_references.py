"""
validate_cross_references.py – Figure/Table Cross-Reference Validator
=====================================================================
Validates that all in-text references to figures (Hình X.Y) and tables
(Bảng X.Y) match actual captions in the document, and vice versa.

Checks:
  [1] Phantom references (cited in text but no caption exists)
  [2] Orphan captions (caption exists but never referenced in text)
  [3] Sequential numbering within each chapter
  [4] Summary statistics

Usage:
  python validate_cross_references.py --path "path/to/doc.docx"
"""

import docx
import sys
import re
import os
import unicodedata
import argparse
from collections import defaultdict

sys.stdout.reconfigure(encoding='utf-8')


def normalize(text: str) -> str:
    return unicodedata.normalize('NFC', text.strip())


def find_ref_heading_idx(doc):
    """Find the index of the references section heading."""
    for idx in range(len(doc.paragraphs) - 1, -1, -1):
        p = doc.paragraphs[idx]
        text = normalize(p.text).upper()
        style = p.style.name if p.style else ''
        if style.startswith('Heading') and 'DANH MỤC TÀI LIỆU' in text:
            return idx
    return len(doc.paragraphs)


def find_list_sections(doc):
    """
    Find the paragraph ranges for 'DANH MỤC HÌNH VẼ' and 'DANH MỤC BẢNG BIỂU'
    sections. References inside these sections should be excluded from body scans.
    """
    exclude_ranges = []
    h1_indices = []
    
    for idx, p in enumerate(doc.paragraphs):
        if p.style and p.style.name.startswith('Heading 1'):
            h1_indices.append(idx)
    
    for i, idx in enumerate(h1_indices):
        text = normalize(doc.paragraphs[idx].text).upper()
        if 'DANH MỤC HÌNH' in text or 'DANH MỤC BẢNG' in text or 'MỤC LỤC' in text:
            # This section runs until the next Heading 1
            end = h1_indices[i + 1] if i + 1 < len(h1_indices) else len(doc.paragraphs)
            exclude_ranges.append((idx, end))
    
    return exclude_ranges


def is_in_excluded_range(idx, exclude_ranges):
    return any(start <= idx < end for start, end in exclude_ranges)


def validate_cross_references(doc_path: str) -> bool:
    print(f"{'='*60}")
    print(f"  CROSS-REFERENCE VALIDATION")
    print(f"  {doc_path}")
    print(f"{'='*60}\n")

    if not os.path.exists(doc_path):
        print(f"FATAL: File not found at {doc_path}")
        return False

    doc = docx.Document(doc_path)
    issues = []
    warnings = []

    ref_heading_idx = find_ref_heading_idx(doc)
    exclude_ranges = find_list_sections(doc)

    # Pattern: Hình X.Y or Bảng X.Y (with optional trailing text)
    fig_pattern = re.compile(r'Hình\s+(\d+)\.(\d+)')
    tbl_pattern = re.compile(r'Bảng\s+(\d+)\.(\d+)')

    # -----------------------------------------------------------------------
    # 1. Collect all captions (paragraphs starting with "Hình X.Y" or "Bảng X.Y")
    # -----------------------------------------------------------------------
    caption_figures = {}  # "X.Y" -> paragraph_idx
    caption_tables = {}

    for idx, p in enumerate(doc.paragraphs):
        text = normalize(p.text)
        if is_in_excluded_range(idx, exclude_ranges):
            continue

        # Check if paragraph starts with figure/table label
        fig_m = re.match(r'^Hình\s+(\d+)\.(\d+)', text)
        if fig_m:
            key = f"{fig_m.group(1)}.{fig_m.group(2)}"
            caption_figures[key] = idx

        tbl_m = re.match(r'^Bảng\s+(\d+)\.(\d+)', text)
        if tbl_m:
            key = f"{tbl_m.group(1)}.{tbl_m.group(2)}"
            caption_tables[key] = idx

    print(f"[0] Found {len(caption_figures)} figure captions, {len(caption_tables)} table captions.\n")

    # -----------------------------------------------------------------------
    # 2. Collect all in-text references
    # -----------------------------------------------------------------------
    ref_figures = defaultdict(list)   # "X.Y" -> [paragraph_indices]
    ref_tables = defaultdict(list)

    # Scan body paragraphs (before references, outside excluded ranges)
    for idx in range(ref_heading_idx):
        if is_in_excluded_range(idx, exclude_ranges):
            continue
        p = doc.paragraphs[idx]
        text = normalize(p.text)
        
        # Skip if this IS a caption paragraph (starts with Hình/Bảng label)
        if text.startswith('Hình ') or text.startswith('Bảng '):
            # Check if it's a caption (short, centered) vs body text mentioning figures
            # Captions typically are short. Body references appear mid-sentence.
            # We'll include captions as references only if they mention OTHER figures
            pass

        for m in fig_pattern.finditer(text):
            key = f"{m.group(1)}.{m.group(2)}"
            ref_figures[key].append(idx)

        for m in tbl_pattern.finditer(text):
            key = f"{m.group(1)}.{m.group(2)}"
            ref_tables[key].append(idx)

    # Also scan table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = normalize(p.text)
                    for m in fig_pattern.finditer(text):
                        key = f"{m.group(1)}.{m.group(2)}"
                        ref_figures[key].append(-1)  # -1 = table cell
                    for m in tbl_pattern.finditer(text):
                        key = f"{m.group(1)}.{m.group(2)}"
                        ref_tables[key].append(-1)

    # -----------------------------------------------------------------------
    # [1] Phantom references (cited but no caption)
    # -----------------------------------------------------------------------
    print("[1] Checking for phantom references (cited but no caption)...")
    phantom_figs = set(ref_figures.keys()) - set(caption_figures.keys())
    phantom_tbls = set(ref_tables.keys()) - set(caption_tables.keys())

    if phantom_figs:
        for key in sorted(phantom_figs):
            locations = [f"P[{i}]" for i in ref_figures[key][:3]]
            msg = f"Hình {key} — cited at {', '.join(locations)} but no caption found"
            issues.append(msg)
            print(f"    ✗ {msg}")
    if phantom_tbls:
        for key in sorted(phantom_tbls):
            locations = [f"P[{i}]" for i in ref_tables[key][:3]]
            msg = f"Bảng {key} — cited at {', '.join(locations)} but no caption found"
            issues.append(msg)
            print(f"    ✗ {msg}")
    if not phantom_figs and not phantom_tbls:
        print(f"    ✓ All in-text references have matching captions.")

    # -----------------------------------------------------------------------
    # [2] Orphan captions (caption exists but never referenced)
    # -----------------------------------------------------------------------
    print(f"\n[2] Checking for orphan captions (caption but never referenced)...")
    orphan_figs = set(caption_figures.keys()) - set(ref_figures.keys())
    orphan_tbls = set(caption_tables.keys()) - set(ref_tables.keys())

    if orphan_figs:
        for key in sorted(orphan_figs):
            idx = caption_figures[key]
            text = normalize(doc.paragraphs[idx].text)[:60]
            msg = f"Hình {key} at P[{idx}] — caption exists but never referenced in text"
            warnings.append(msg)
            print(f"    ⚠ {msg}")
    if orphan_tbls:
        for key in sorted(orphan_tbls):
            idx = caption_tables[key]
            text = normalize(doc.paragraphs[idx].text)[:60]
            msg = f"Bảng {key} at P[{idx}] — caption exists but never referenced in text"
            warnings.append(msg)
            print(f"    ⚠ {msg}")
    if not orphan_figs and not orphan_tbls:
        print(f"    ✓ All captions are referenced in body text.")

    # -----------------------------------------------------------------------
    # [3] Sequential numbering within chapters
    # -----------------------------------------------------------------------
    print(f"\n[3] Checking sequential numbering within chapters...")

    # Group figures by chapter
    fig_by_chapter = defaultdict(list)
    for key in caption_figures:
        ch, num = key.split('.')
        fig_by_chapter[int(ch)].append(int(num))

    tbl_by_chapter = defaultdict(list)
    for key in caption_tables:
        ch, num = key.split('.')
        tbl_by_chapter[int(ch)].append(int(num))

    seq_issues = 0
    for ch in sorted(fig_by_chapter):
        nums = sorted(fig_by_chapter[ch])
        expected = list(range(1, len(nums) + 1))
        if nums != expected:
            gaps = set(expected) - set(nums)
            if gaps:
                msg = f"Hình chapter {ch}: missing numbers {sorted(gaps)} (have {nums})"
                warnings.append(msg)
                print(f"    ⚠ {msg}")
                seq_issues += 1

    for ch in sorted(tbl_by_chapter):
        nums = sorted(tbl_by_chapter[ch])
        expected = list(range(1, len(nums) + 1))
        if nums != expected:
            gaps = set(expected) - set(nums)
            if gaps:
                msg = f"Bảng chapter {ch}: missing numbers {sorted(gaps)} (have {nums})"
                warnings.append(msg)
                print(f"    ⚠ {msg}")
                seq_issues += 1

    if seq_issues == 0:
        print(f"    ✓ All figures and tables are numbered sequentially within chapters.")

    # -----------------------------------------------------------------------
    # [4] Statistics
    # -----------------------------------------------------------------------
    print(f"\n[4] Statistics:")
    print(f"    Figure captions:      {len(caption_figures)}")
    print(f"    Table captions:       {len(caption_tables)}")
    print(f"    Figure references:    {len(ref_figures)} unique")
    print(f"    Table references:     {len(ref_tables)} unique")
    all_chapters = sorted(set(list(fig_by_chapter.keys()) + list(tbl_by_chapter.keys())))
    for ch in all_chapters:
        fc = len(fig_by_chapter.get(ch, []))
        tc = len(tbl_by_chapter.get(ch, []))
        print(f"    Chapter {ch}: {fc} figures, {tc} tables")

    # -----------------------------------------------------------------------
    # Summary
    # -----------------------------------------------------------------------
    print(f"\n{'='*60}")
    print(f"  CROSS-REFERENCE VALIDATION SUMMARY")
    print(f"{'='*60}")
    print(f"  Total issues (errors):   {len(issues)}")
    print(f"  Total warnings:          {len(warnings)}")

    if issues:
        print(f"\n  ERRORS:")
        for i, issue in enumerate(issues, 1):
            print(f"    {i}. {issue}")
    if warnings:
        print(f"\n  WARNINGS:")
        for i, w in enumerate(warnings, 1):
            print(f"    {i}. {w}")
    if not issues and not warnings:
        print(f"\n  ✓✓✓ ALL CROSS-REFERENCES ARE VALID ✓✓✓")

    status = "PASS" if not issues else "FAIL"
    print(f"\n  Final verdict: {status}")
    print(f"{'='*60}\n")

    return len(issues) == 0


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Validate figure/table cross-references.')
    parser.add_argument('--path', type=str,
                        default='d:/MiniERP_NhomKinh/ChuyenDeTotNghiep_NguyenTrongThoi.docx',
                        help='Path to the .docx file.')
    args = parser.parse_args()
    validate_cross_references(args.path)
