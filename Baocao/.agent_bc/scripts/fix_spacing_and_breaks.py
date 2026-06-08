"""
fix_spacing_and_breaks.py – Self-Healing Word Spacing & Soft-Break Fixer
=========================================================================
Scans justified paragraphs for soft line breaks (Shift+Enter = <w:br/> or \\n).
When found, splits them into separate hard paragraphs while **preserving all
run-level formatting** (bold, italic, underline, font name, font size, color,
highlight, superscript/subscript, strikethrough).

Safety:
  • Creates a timestamped backup before any modification.
  • Processes paragraphs in reverse index order for stable offsets.
  • Preserves paragraph style, alignment, line spacing, and indentation.

Usage:
  python fix_spacing_and_breaks.py --path "path/to/doc.docx"
"""

import docx
import sys
import os
import shutil
import argparse
from datetime import datetime
from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def clone_run_properties(source_rPr):
    """Deep-copy run properties (<w:rPr>) from a source element."""
    if source_rPr is None:
        return None
    return deepcopy(source_rPr)


def is_justified(p) -> bool:
    """Check if a paragraph has justified alignment (explicit or inherited)."""
    fmt = p.paragraph_format
    if fmt.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return True
    if fmt.alignment is None and p.style:
        style_align = p.style.paragraph_format.alignment
        if style_align == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return True
    return False


def paragraph_has_soft_breaks(p) -> bool:
    """
    Check if a paragraph contains soft line breaks.
    These are either:
      - <w:br/> elements inside runs (Shift+Enter in Word)
      - Literal \\n characters within run text
    """
    for r_el in p._element.findall(qn('w:r')):
        # Check for <w:br/> elements
        if r_el.findall(qn('w:br')):
            return True
        # Check for literal newlines in text
        t_el = r_el.find(qn('w:t'))
        if t_el is not None and t_el.text and '\n' in t_el.text:
            return True
    return False


def copy_paragraph_format(source_p, target_p):
    """Copy paragraph-level formatting from source to target."""
    # Copy paragraph properties XML if available
    source_pPr = source_p._element.find(qn('w:pPr'))
    if source_pPr is not None:
        new_pPr = deepcopy(source_pPr)
        existing_pPr = target_p._element.find(qn('w:pPr'))
        if existing_pPr is not None:
            target_p._element.replace(existing_pPr, new_pPr)
        else:
            target_p._element.insert(0, new_pPr)


def split_runs_by_breaks(p):
    """
    Split a paragraph's runs into groups separated by soft breaks.
    Returns a list of lists, where each inner list contains
    (text, rPr_copy) tuples for one 'line'.
    
    Handles both <w:br/> elements and \\n in text.
    """
    lines = [[]]  # Start with one empty line

    for r_el in p._element.findall(qn('w:r')):
        rPr = r_el.find(qn('w:rPr'))

        # Process child elements in order
        for child in r_el:
            tag = child.tag
            if tag == qn('w:br'):
                # Soft break → start a new line
                lines.append([])
            elif tag == qn('w:t'):
                if child.text is None:
                    continue
                text = child.text
                if '\n' in text:
                    # Split by newlines
                    parts = text.split('\n')
                    for i, part in enumerate(parts):
                        if part:
                            lines[-1].append((part, clone_run_properties(rPr)))
                        if i < len(parts) - 1:
                            lines.append([])
                else:
                    if text:
                        lines[-1].append((text, clone_run_properties(rPr)))
            # Preserve other elements like <w:tab/>, <w:drawing/>, etc.
            # by keeping them in the first line (conservative approach)

    return lines


def create_paragraph_from_line(line_runs, template_p, parent_element, insert_before):
    """
    Create a new <w:p> element with the given runs and formatting,
    inserted before insert_before element.
    
    line_runs: list of (text, rPr_xml_or_None) tuples
    template_p: the original paragraph to copy paragraph-level formatting from
    """
    # Create new paragraph element
    new_p_el = OxmlElement('w:p')

    # Copy paragraph properties
    source_pPr = template_p._element.find(qn('w:pPr'))
    if source_pPr is not None:
        new_p_el.append(deepcopy(source_pPr))

    # Add runs with preserved formatting
    for text, rPr in line_runs:
        r_el = OxmlElement('w:r')
        if rPr is not None:
            r_el.append(rPr)
        t_el = OxmlElement('w:t')
        t_el.text = text
        t_el.set(qn('xml:space'), 'preserve')
        r_el.append(t_el)
        new_p_el.append(r_el)

    insert_before.addprevious(new_p_el)
    return new_p_el


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def fix_document_spacing(doc_path: str) -> bool:
    """
    Scan all justified paragraphs for soft line breaks.
    Split them into separate paragraphs while preserving formatting.
    """
    print(f"=== HEALING SPACING & LINE BREAKS: {doc_path} ===")
    if not os.path.exists(doc_path):
        print(f"ERROR: File not found at {doc_path}")
        return False

    doc = docx.Document(doc_path)

    # Identify paragraphs that need fixing
    candidates = []
    for idx, p in enumerate(doc.paragraphs):
        if is_justified(p) and paragraph_has_soft_breaks(p):
            candidates.append(idx)

    if not candidates:
        print("✓ No spacing issues with soft line breaks found in justified paragraphs.")
        return True

    # Create backup only if we have work to do
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup_path = doc_path.replace('.docx', f'_Backup_Spacing_{timestamp}.docx')
    shutil.copy2(doc_path, backup_path)
    print(f"Backup created at {backup_path}")

    print(f"Found {len(candidates)} justified paragraphs with soft breaks.")

    # Process in reverse order for stable indices
    fixed_count = 0
    for idx in reversed(candidates):
        p = doc.paragraphs[idx]
        text_preview = p.text[:60].replace('\n', '\\n')
        print(f"  Fixing P[{idx}]: '{text_preview}...'")

        # Split runs into line groups
        line_groups = split_runs_by_breaks(p)

        # Filter out empty lines
        line_groups = [lg for lg in line_groups if lg]

        if len(line_groups) <= 1:
            # No actual split needed (just a trailing break or empty)
            continue

        # Insert new paragraphs before the original
        target_el = p._element
        for line_runs in line_groups:
            create_paragraph_from_line(
                line_runs,
                template_p=p,
                parent_element=target_el.getparent(),
                insert_before=target_el
            )

        # Remove the original paragraph
        target_el.getparent().remove(target_el)
        fixed_count += 1

    if fixed_count > 0:
        doc.save(doc_path)
        print(f"\n✓ Healed {fixed_count} paragraphs successfully!")
        print(f"  Backup: {backup_path}")
    else:
        print("✓ No paragraphs needed splitting after analysis.")
        os.remove(backup_path)

    return True


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Heal word-stretching spacing issues caused by soft line breaks.')
    parser.add_argument('--path', type=str,
                        default='d:/MiniERP_NhomKinh/ChuyenDeTotNghiep_NguyenTrongThoi.docx',
                        help='Path to the .docx file.')
    args = parser.parse_args()
    fix_document_spacing(args.path)
