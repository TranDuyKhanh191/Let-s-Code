import docx
import re
import os

def markdown_to_docx(md_filepath, docx_filepath):
    doc = docx.Document()
    
    try:
        with open(md_filepath, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('---'):
                    # Horizontal rule approximation
                    doc.add_paragraph('_' * 40)
                else:
                    # check if list
                    p = None
                    if line.startswith('- ') or line.startswith('* '):
                        p = doc.add_paragraph(style='List Bullet')
                        line = line[2:]
                    else:
                        p = doc.add_paragraph()
                    
                    # Process bold (**text**) and italic (*text*)
                    # Simple regex parsing
                    parts = re.split(r'(\*\*.*?\*\*|\*[^\*]+\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            p.add_run(part[2:-2]).bold = True
                        elif part.startswith('*') and part.endswith('*'):
                            p.add_run(part[1:-1]).italic = True
                        else:
                            p.add_run(part)
                            
        doc.save(docx_filepath)
        print(f"Successfully created {docx_filepath}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    md_file = "CauTrucMoi_Chuong2_3_4.md"
    docx_file = "CauTrucMoi_Chuong2_3_4.docx"
    markdown_to_docx(md_file, docx_file)
