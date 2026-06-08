import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_style(run, bold=False, italic=False, size=13):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_style(run, bold=True, size=15)
    else:
        apply_style(run, bold=True, italic=(level==3), size=13)

def add_paragraph(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(36)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    apply_style(run, size=13, italic=italic, bold=bold)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    apply_style(run, size=13)

def create_chuong4():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    add_heading(doc, "CHƯƠNG IV. QUY TRÌNH XỬ LÝ VÀ KẾT QUẢ ĐẠT ĐƯỢC", 0)
    add_paragraph(doc, "Chương này đi sâu vào luồng đi của dữ liệu (Data Flow), cách ứng dụng Client quản lý trạng thái, tương tác với cơ sở dữ liệu thông qua các phương thức RESTful API, cũng như trình bày kết quả hiện thực hóa các chức năng từ cơ bản đến nâng cao (Giai đoạn 2) của nền tảng.")

    add_heading(doc, "4.1. QUY TRÌNH XÂY DỰNG VÀ TRIỂN KHAI CƠ SỞ DỮ LIỆU TRÊN SUPABASE", 1)
    
    add_heading(doc, "4.1.1. Lý do lựa chọn và Triển khai cơ sở dữ liệu", 2)
    add_paragraph(doc, "Supabase được xây dựng trên nền tảng PostgreSQL, cung cấp giao diện quản trị trực quan, tự động sinh ra các API truy xuất dữ liệu theo chuẩn REST và hỗ trợ bảo mật mạnh mẽ. Quy trình triển khai bao gồm: phân tích yêu cầu từ phòng Giáo vụ, xây dựng mô hình dữ liệu logic/vật lý, và thực thi các lệnh SQL để tạo toàn bộ bảng dữ liệu. Toàn bộ lược đồ được kiểm tra lại để tránh trùng lặp và sai ràng buộc.")
    
    add_heading(doc, "4.1.2. Quy trình nhập dữ liệu thử nghiệm", 2)
    add_paragraph(doc, "Trong giai đoạn thử nghiệm, dữ liệu được lấy từ tài liệu bài học Robotics thực tế (ví dụ: mô hình Nữ Vũ Công). Quy trình nhập liệu bao gồm việc xác định khóa học, phân rã mục tiêu (kiến thức, kỹ năng, thái độ), nhập nội dung lý thuyết, thêm hoạt động thử thách, câu hỏi trắc nghiệm và cuối cùng là gắn tài nguyên đa phương tiện. Việc này kiểm tra tính hợp lý của quan hệ bảng và khả năng đáp ứng thực tế.")

    add_heading(doc, "4.2. KIẾN TRÚC BẢO MẬT VÀ AN TOÀN DỮ LIỆU", 1)

    add_heading(doc, "4.2.1. Phân quyền và Bảo mật cơ bản", 2)
    add_paragraph(doc, "Hệ thống áp dụng các nguyên tắc: phân quyền truy cập theo vai trò (admin, teacher, student), tối thiểu hóa quyền hạn, đảm bảo toàn vẹn dữ liệu (khóa ngoại, kiểu dữ liệu), và truy vết lịch sử thay đổi (bảng assignment_logs).")

    add_heading(doc, "4.2.2. Row Level Security (RLS) ở mức Database", 2)
    add_paragraph(doc, "Ở Giai đoạn 2, tính năng RLS (Row Level Security) được kích hoạt trên Supabase. Thông qua các chính sách (Policies), dữ liệu được bảo vệ tận gốc ở tầng DB. Ví dụ: Dù người dùng có cố gắng gọi API trực tiếp, hệ thống vẫn dựa vào Token định danh để chặn học sinh đọc dữ liệu khóa học không thuộc về mình, hoặc giáo viên chỉ xem được assignment của chính mình. Điều này ngăn chặn hoàn toàn nguy cơ truy cập trái phép qua Client SDK.")

    add_heading(doc, "4.2.3. Cơ chế Auth và Quản lý phiên làm việc", 2)
    add_paragraph(doc, "Hệ thống thiết lập quy trình xác thực chặt chẽ dựa trên JSON Web Token (JWT). JWT được tự động lưu trữ vào LocalStorage và đính kèm (Request Interceptor) vào Header Authorization của mọi gói tin HTTP. Điều này đảm bảo trạng thái đăng nhập được duy trì mượt mà.")

    add_heading(doc, "4.3. QUY TRÌNH XỬ LÝ DỮ LIỆU VÀ TƯƠNG TÁC API", 1)

    add_heading(doc, "4.3.1. Kỹ thuật Fetching dữ liệu song song (Promise.all)", 2)
    add_paragraph(doc, "Trong module quản lý khóa học, dữ liệu phân tán ở nhiều endpoint. Để tối ưu hóa, hệ thống sử dụng Promise.all để gọi API song song (Non-blocking) thay vì tuần tự, sau đó gộp (Merge) và sắp xếp (Sort) ngay tại Client để hiển thị nhanh chóng.")

    add_heading(doc, "4.3.2. Logic \"Upsert\" và Xử lý hàng loạt", 2)
    add_paragraph(doc, "Tại các form như LessonObjectives, hệ thống tự động kiểm tra ID: nếu ID tồn tại sẽ gửi PATCH (Cập nhật), nếu ID null sẽ gửi POST (Thêm mới). Đối với thao tác xử lý hàng loạt như Duyệt phân công nhanh, hệ thống tạo ra một mảng các Promise và gửi đi cùng lúc, giúp cập nhật trạng thái của hàng chục bản ghi chỉ trong chớp mắt.")

    add_heading(doc, "4.3.3. Tối ưu hóa API Aggregation bằng PostgreSQL Function", 2)
    add_paragraph(doc, "Để giải quyết bài toán tải chậm do gọi nhiều API riêng lẻ, hệ thống xây dựng một hàm PostgreSQL get_aggregated_lesson để gộp nội dung từ 6 bảng (lessons, objectives, contents, quizzes, challenges, media) thành một JSON payload duy nhất bằng json_agg. Lớp Service Layer của hệ thống chỉ cần gọi hàm supabase.rpc(), giúp giảm độ trễ API xuống dưới 100ms.")

    add_heading(doc, "4.4. XỬ LÝ LOGIC QUAN HỆ VÀ ĐA PHƯƠNG TIỆN", 1)

    add_heading(doc, "4.4.1. Lọc quan hệ Nhiều-Nhiều (M-N) và Xóa liên hoàn", 2)
    add_paragraph(doc, "Trong chức năng phân công, hệ thống sử dụng thuật toán lọc mảng để so sánh khóa học đang chọn với lịch sử phân công. Các khóa học đã gán sẽ tự động bị vô hiệu hóa (Disabled), ngăn chặn lỗi ràng buộc duy nhất (Unique Constraint) từ phía Client. Khi xóa dữ liệu (ví dụ: xóa Thử thách), hệ thống mô phỏng Cascade Delete: xóa liên kết trước, xóa file vật lý trên Storage sau, và cuối cùng mới xóa bản ghi cha để tránh để lại file rác.")

    add_heading(doc, "4.4.2. Quy trình Upload và Liên kết File Đa phương tiện", 2)
    add_paragraph(doc, "Xử lý Media trải qua quy trình 2 bước (Atomic Upload Transaction): Bước 1 gửi file nhị phân qua FormData lên Storage để nhận mediaId. Bước 2 dùng mediaId tạo bản ghi trong bảng trung gian kèm phân loại purpose (ví dụ: 'intro', 'gallery', 'main') để Component dễ dàng lọc và hiển thị đúng chỗ.")

    add_heading(doc, "4.4.3. Xử lý Dữ liệu phức tạp và Lồng nhau (Nested JSON)", 2)
    add_paragraph(doc, "Với chức năng Trắc nghiệm (Quiz), Frontend đóng gói toàn bộ danh sách câu hỏi và phương án vào một cấu trúc JSON lồng nhau (Nested JSON Payload). Điều này giúp chỉ tốn 1 round-trip gửi xuống Server thay vì nhiều request rời rạc.")

    add_heading(doc, "4.5. HIỆN THỰC HÓA CÁC TÍNH NĂNG ĐỘT PHÁ (GIAI ĐOẠN 2)", 1)

    add_heading(doc, "4.5.1. Cổng kết nối Lego SPIKE và Nộp bài (.llsp3)", 2)
    add_paragraph(doc, "Hệ thống kết nối luồng học tập với môi trường lập trình của Lego Education thông qua Deep Link (Mở không gian Lego). Đặc biệt, hệ thống xử lý thao tác nộp bài bằng Modal Drag & Drop cho phép học sinh kéo thả file dự án (.llsp3). File được upload multipart và lưu trữ vào phân vùng riêng biệt trong Supabase Storage.")

    add_heading(doc, "4.5.2. Trợ lý AI Phân tích Hình ảnh (AI Vision Tutor) với SSE", 2)
    add_paragraph(doc, "Tích hợp mô hình Gemini 1.5 Pro Vision đóng vai trò gia sư AI. Học sinh chụp ảnh đoạn khối lệnh lỗi và gửi vào chat. AI phân tích logic và đưa ra gợi ý sư phạm (không đưa trực tiếp đáp án). Để mang lại cảm giác thân thiện như ChatGPT, luồng giao tiếp sử dụng công nghệ Server-Sent Events (SSE), cho phép chữ của AI được stream (chữ chạy liên tục) về Frontend mà không bị trễ.")

    add_heading(doc, "4.5.3. Gamification và Hiệu ứng khen thưởng trực quan", 2)
    add_paragraph(doc, "Tích hợp thư viện canvas-confetti kết hợp GSAP. Khi API trả về HTTP 200 (Nộp bài thành công), màn hình lập tức bắn pháo hoa từ nút bấm, mang lại hiệu ứng thị giác bùng nổ (Wow Factor) và khích lệ mạnh mẽ động lực học tập của học sinh.")

    add_heading(doc, "4.5.4. Thông báo Real-time Sync", 2)
    add_paragraph(doc, "Sử dụng tính năng Supabase Realtime, hệ thống thiết lập WebSocket lắng nghe sự thay đổi trên bảng phân công. Khi Admin gán lớp học mới, thông báo (Toast) lập tức xuất hiện trên màn hình giáo viên và danh sách tự làm mới mà không cần thao tác F5.")

    add_heading(doc, "4.6. KẾT QUẢ ĐẠT ĐƯỢC VÀ DEMO SẢN PHẨM", 1)
    
    add_heading(doc, "4.6.1. Giao diện Đăng nhập và Xác thực", 2)
    add_paragraph(doc, "Hệ thống cung cấp giao diện đăng nhập thống nhất cho cả 3 vai trò: Quản trị viên, Giáo viên và Học sinh.")
    add_paragraph(doc, "(Chèn ảnh: Giao diện Đăng nhập hệ thống)", italic=True)
    add_paragraph(doc, "Hình 4.1: Giao diện Đăng nhập hệ thống (Login)", italic=True, bold=True)

    add_heading(doc, "4.6.2. Phân hệ Quản trị (Admin Portal)", 2)
    add_paragraph(doc, "Giao diện quản trị hiện đại theo phong cách Glassmorphism. Dashboard cung cấp cái nhìn toàn cảnh qua các thẻ thống kê (Stats Cards) và lịch trình (Calendar Widget).")
    add_paragraph(doc, "(Chèn ảnh: Dashboard Tổng quan Admin)", italic=True)
    add_paragraph(doc, "Hình 4.2: Giao diện Dashboard Tổng quan của Quản trị viên", italic=True, bold=True)

    add_paragraph(doc, "CMS Bài giảng xử lý mượt mà quá trình soạn thảo đa phương tiện. Admin có thể xem danh sách, thêm mới khóa học, bài học và chỉnh sửa chi tiết từng bài học thông qua giao diện kéo thả trực quan.")
    add_paragraph(doc, "(Chèn ảnh: Giao diện thêm mới và chỉnh sửa Bài học)", italic=True)
    add_paragraph(doc, "Hình 4.3: Giao diện thêm mới và chỉnh sửa chi tiết Bài học", italic=True, bold=True)
    
    add_paragraph(doc, "Quản lý nhân sự: Hệ thống cho phép thêm mới, chỉnh sửa thông tin giáo viên và học sinh một cách nhanh chóng, đảm bảo tính bảo mật và chuẩn xác.")
    add_paragraph(doc, "(Chèn ảnh: Giao diện danh sách và chỉnh sửa thông tin Giáo viên/Học sinh)", italic=True)
    add_paragraph(doc, "Hình 4.4: Giao diện quản lý và thêm mới Giáo viên/Học sinh", italic=True, bold=True)

    add_paragraph(doc, "Hệ thống AssignmentModal điều phối nhân sự thông minh, chống trùng lặp dữ liệu phân công.")
    add_paragraph(doc, "(Chèn ảnh: Modal phân công lớp học)", italic=True)
    add_paragraph(doc, "Hình 4.5: Giao diện Modal phân công Khóa học cho Giáo viên", italic=True, bold=True)

    add_heading(doc, "4.6.3. Phân hệ Cổng thông tin Giáo viên (Teacher Portal)", 2)
    add_paragraph(doc, "Cung cấp lưới khóa học trực quan với bộ lọc Real-time. Điểm nhấn là Trình phát bài giảng (Lesson Player) áp dụng mô hình StageRenderer, chuyển đổi liền mạch giữa Slide, Video và Quiz mà không tải lại trang. Các cơ chế như Smart Pre-fetching (tải trước ảnh) và ghi nhớ vị trí Video (savedVideoTime) giúp thao tác dạy học tại lớp trơn tru tuyệt đối.")
    add_paragraph(doc, "(Chèn ảnh: Lưới khóa học của giáo viên)", italic=True)
    add_paragraph(doc, "Hình 4.6: Giao diện danh sách khóa học được phân công của Giáo viên", italic=True, bold=True)

    add_paragraph(doc, "Giao diện theo dõi lớp học cho phép giáo viên thấy ngay tiến độ nộp bài (.llsp3) của từng học sinh, tích hợp tính năng tải file bài làm chỉ với 1 click.")
    add_paragraph(doc, "(Chèn ảnh: Giao diện theo dõi tiến độ nộp bài của học sinh)", italic=True)
    add_paragraph(doc, "Hình 4.7: Giao diện theo dõi học sinh và Tải file bài nộp Lego SPIKE", italic=True, bold=True)

    add_heading(doc, "4.6.4. Phân hệ Cổng Tự học Học sinh (Student Portal)", 2)
    add_paragraph(doc, "Giao diện được tinh giản, tái sử dụng nội dung (Component Reuse) từ bài giảng của giáo viên nhưng ẩn đi các công cụ quản trị. Học sinh được trải nghiệm lộ trình tự học liền mạch: xem lý thuyết, thực hành trên nền tảng Lego SPIKE, kéo thả file bài làm (.llsp3) để nộp, nhận pháo hoa chúc mừng và tương tác với Trợ lý AI khi gặp lỗi.")
    add_paragraph(doc, "(Chèn ảnh: Giao diện học tập và nút kéo thả nộp bài)", italic=True)
    add_paragraph(doc, "Hình 4.8: Không gian bài học của Học sinh và giao diện nộp bài tập", italic=True, bold=True)
    
    add_paragraph(doc, "(Chèn ảnh: Khung chat AI Tutor giải đáp lỗi khối lệnh)", italic=True)
    add_paragraph(doc, "Hình 4.9: Giao diện tương tác với Trợ lý AI Vision Tutor và hiệu ứng Gamification", italic=True, bold=True)

    output_path = os.path.join(os.path.dirname(__file__), 'Chuong4_HienThucHoa.docx')
    doc.save(output_path)
    print(f"File created successfully at {output_path}")

if __name__ == '__main__':
    create_chuong4()
