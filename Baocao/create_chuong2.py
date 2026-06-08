from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def set_font(run, bold=False, italic=False, size=13):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_chapter(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, size=15)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, size=13)

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, size=13)

def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, italic=True, size=13)

def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=13)
    
def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=13)

def generate_document(filepath):
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    add_chapter(doc, "Chương 2. CƠ SỞ LÝ THUYẾT")
    
    add_heading1(doc, "2.1 TỔNG QUAN VỀ HỆ THỐNG QUẢN LÝ HỌC TẬP VÀ CÔNG NGHỆ GIÁO DỤC")

    add_heading2(doc, "2.1.1 Khái niệm EdTech và LMS")
    add_para(doc, "Chuyển đổi số trong giáo dục (EdTech - Educational Technology) là việc ứng dụng các công cụ phần mềm, phần cứng và các quy trình công nghệ hiện đại vào môi trường lớp học để nâng cao hiệu quả giảng dạy và học tập. EdTech không chỉ dừng lại ở việc đưa thiết bị điện tử vào lớp học mà còn bao gồm việc tạo ra các phương thức tương tác mới giữa người dạy và người học.")
    add_para(doc, "Hệ thống quản lý học tập (LMS - Learning Management System) là một ứng dụng phần mềm dùng để quản trị, lập tài liệu, theo dõi, báo cáo và phân phối các chương trình giáo dục. LMS cung cấp nền tảng trực tuyến giúp giáo viên có thể thiết kế bài giảng, giao bài tập, đánh giá năng lực học sinh. Đồng thời, hệ thống quản trị nội dung (CMS - Content Management System) thường được tích hợp cùng LMS để quản lý vòng đời của các tài nguyên số (hình ảnh, video, tài liệu văn bản) được sử dụng trong khóa học.")

    add_heading2(doc, "2.1.2 Lý thuyết về phương pháp sư phạm cho Robotics")
    add_para(doc, "Đào tạo STEM nói chung và Robotics nói riêng đòi hỏi một phương pháp sư phạm thực hành đặc thù, kết hợp giữa tư duy logic và kỹ năng thao tác. Một trong những quy trình giảng dạy tiêu chuẩn được áp dụng phổ biến bao gồm 8 bước sư phạm liên hoàn:")
    add_bullet(doc, "Mục tiêu (Objective): Xác định rõ chuẩn đầu ra của bài học, những gì người học cần đạt được về mặt kiến thức và kỹ năng sau khi hoàn thành.")
    add_bullet(doc, "Tổng quan (Overview): Cung cấp cái nhìn bao quát về bối cảnh của bài học, kích thích sự tò mò của người học.")
    add_bullet(doc, "Giới thiệu (Introduction): Diễn giải các khái niệm mới, nguyên lý hoạt động của các thiết bị cơ khí hoặc điện tử sắp được sử dụng.")
    add_bullet(doc, "Chuẩn bị (Preparation): Liệt kê các công cụ, linh kiện, cảm biến và phần mềm cần thiết trước khi bắt tay vào thực hành.")
    add_bullet(doc, "Xây dựng (Build): Hướng dẫn chi tiết từng bước (step-by-step) cách lắp ráp mô hình robot vật lý hoặc thiết lập mạch điện.")
    add_bullet(doc, "Lý thuyết lập trình (Theory): Phân tích nguyên lý của các khối lệnh, cấu trúc điều kiện hoặc vòng lặp sẽ được sử dụng để điều khiển mô hình.")
    add_bullet(doc, "Thử thách (Challenge): Đưa ra các bài toán mở yêu cầu người học vận dụng kiến thức vừa học để tùy biến hoặc nâng cấp hoạt động của robot.")
    add_bullet(doc, "Tổng kết (Summary): Đánh giá lại kết quả học tập, hệ thống hóa kiến thức và rút ra kinh nghiệm.")

    add_heading1(doc, "2.2 CƠ SỞ LÝ THUYẾT VỀ KIẾN TRÚC PHẦN MỀM VÀ CƠ SỞ DỮ LIỆU")

    add_heading2(doc, "2.2.1 Cơ sở dữ liệu quan hệ và Chuẩn hóa dữ liệu")
    add_para(doc, "Hệ quản trị cơ sở dữ liệu quan hệ (RDBMS - Relational Database Management System) là hệ thống lưu trữ dữ liệu dưới dạng các bảng (tables) bao gồm hàng và cột. Các bảng này có mối quan hệ với nhau thông qua các khóa chính (Primary Key) và khóa ngoại (Foreign Key), đảm bảo tính toàn vẹn dữ liệu (Data Integrity).")
    add_para(doc, "Chuẩn hóa dữ liệu (Data Normalization) là quá trình tổ chức cấu trúc cơ sở dữ liệu nhằm giảm thiểu sự dư thừa và phụ thuộc bất thường. Dạng chuẩn 3 (3NF - Third Normal Form) là mức chuẩn hóa phổ biến nhất trong thiết kế phần mềm. Để đạt được 3NF, cơ sở dữ liệu phải thỏa mãn dạng chuẩn 2 (2NF), đồng thời mọi thuộc tính không khóa phải phụ thuộc trực tiếp vào khóa chính, nghĩa là không được tồn tại sự phụ thuộc bắc cầu (transitive dependency). Việc áp dụng 3NF giúp loại bỏ tối đa việc lặp lại dữ liệu, từ đó giúp tiết kiệm không gian lưu trữ và tránh các lỗi xung đột khi cập nhật dữ liệu.")

    add_heading2(doc, "2.2.2 Kiến trúc phân tầng (Layered Architecture)")
    add_para(doc, "Kiến trúc phân tầng là một mẫu thiết kế phần mềm trong đó các thành phần được tổ chức thành các lớp (layers) nằm chồng lên nhau. Mỗi lớp có một trách nhiệm cụ thể và chỉ giao tiếp với lớp ngay bên dưới nó. Trong phát triển Backend hiện đại, mô hình 4 lớp (Routes - Controllers - Services - Models) thường được sử dụng rộng rãi:")
    add_bullet(doc, "Lớp Routes (Định tuyến): Nhận và phân giải các yêu cầu mạng từ người dùng đến đúng hàm xử lý.")
    add_bullet(doc, "Lớp Controllers (Điều khiển): Đóng vai trò là trạm trung chuyển, tiếp nhận dữ liệu đầu vào, thực hiện xác thực định dạng và gửi kết quả phản hồi.")
    add_bullet(doc, "Lớp Services (Nghiệp vụ): Chứa toàn bộ các quy tắc và logic nghiệp vụ phức tạp của ứng dụng (như tính toán, xử lý ràng buộc). Tách biệt logic nghiệp vụ ra khỏi Controller giúp mã nguồn dễ kiểm thử và tái sử dụng.")
    add_bullet(doc, "Lớp Models (Mô hình dữ liệu): Chịu trách nhiệm tương tác trực tiếp với cơ sở dữ liệu, thực thi các câu lệnh truy vấn (CRUD).")

    add_heading2(doc, "2.2.3 Kiến trúc Component-Based và Trạng thái (State)")
    add_para(doc, "Kiến trúc dựa trên thành phần (Component-Based Architecture) là triết lý thiết kế Frontend hiện đại, trong đó giao diện người dùng được chia nhỏ thành các khối độc lập, có tính đóng gói và có thể tái sử dụng (components). Mỗi component quản lý một phần nhỏ của giao diện và có thể được lắp ghép lại để tạo thành các ứng dụng web phức tạp.")
    add_para(doc, "Nguyên lý luồng dữ liệu một chiều (Unidirectional Data Flow) chỉ ra rằng dữ liệu (State) chỉ di chuyển theo một hướng duy nhất từ component cha xuống component con thông qua các thuộc tính (props). Khi component con muốn thay đổi dữ liệu, nó không được phép sửa trực tiếp mà phải gửi tín hiệu (events) lên component cha để yêu cầu thay đổi. Điều này giúp hệ thống dễ dàng theo dõi sự biến đổi của dữ liệu và hạn chế lỗi xảy ra.")

    add_heading2(doc, "2.2.4 Giao diện lập trình ứng dụng (RESTful API)")
    add_para(doc, "API (Application Programming Interface) là tập hợp các quy tắc và phương thức cho phép các ứng dụng phần mềm khác nhau giao tiếp với nhau. RESTful API là một kiến trúc thiết kế API dựa trên tiêu chuẩn web, sử dụng giao thức HTTP.")
    add_para(doc, "Kiến trúc này sử dụng các phương thức HTTP tiêu chuẩn để định nghĩa loại hành động cần thực hiện đối với tài nguyên:")
    add_bullet(doc, "GET: Truy xuất dữ liệu từ máy chủ.")
    add_bullet(doc, "POST: Tạo mới một tài nguyên trên máy chủ.")
    add_bullet(doc, "PUT: Cập nhật toàn bộ nội dung của một tài nguyên hiện có.")
    add_bullet(doc, "PATCH: Cập nhật một phần của tài nguyên.")
    add_bullet(doc, "DELETE: Xóa một tài nguyên khỏi máy chủ.")

    add_heading1(doc, "2.3 CÁC CÔNG NGHỆ VÀ CÔNG CỤ LẬP TRÌNH ÁP DỤNG")

    add_heading2(doc, "2.3.1 Về phía Backend và Cơ sở dữ liệu")

    add_heading3(doc, "2.3.1.1 Nền tảng Node.js")
    add_para(doc, "Định nghĩa: Node.js là một môi trường thực thi mã nguồn mở đa nền tảng, được xây dựng dựa trên V8 JavaScript engine của Chrome, cho phép lập trình viên chạy mã JavaScript ở phía máy chủ (Server-side).")
    add_para(doc, "Đặc điểm nổi bật: Cơ chế hoạt động bất đồng bộ (Asynchronous) và hướng sự kiện (Event-driven) thông qua Event Loop. Nhờ đó, Node.js không tạo ra các luồng (threads) mới cho mỗi yêu cầu mà xử lý chúng trên một luồng duy nhất (Single-threaded) với mô hình non-blocking I/O.")
    add_para(doc, "Ưu điểm: Tốc độ xử lý nhanh, khả năng chịu tải cao, lý tưởng cho các ứng dụng web thời gian thực hoặc các hệ thống API có cường độ I/O lớn.")
    add_para(doc, "Nhược điểm: Không phù hợp với các tác vụ xử lý tính toán nặng (CPU-intensive) vì mô hình đơn luồng sẽ làm nghẽn toàn bộ hệ thống.")

    add_heading3(doc, "2.3.1.2 Nền tảng Supabase và Hệ quản trị PostgreSQL")
    add_para(doc, "Định nghĩa: Supabase là một nền tảng Backend-as-a-Service (BaaS) mã nguồn mở, cung cấp sẵn các dịch vụ lõi như cơ sở dữ liệu, xác thực người dùng và lưu trữ tệp, giúp lập trình viên không phải tự xây dựng hạ tầng backend từ đầu.")
    add_para(doc, "Đặc điểm nổi bật: Khác với nhiều nền tảng BaaS sử dụng cơ sở dữ liệu NoSQL, Supabase được xây dựng xoay quanh PostgreSQL - một hệ quản trị cơ sở dữ liệu quan hệ mạnh mẽ, chuẩn mực và có khả năng mở rộng rất cao. PostgreSQL hỗ trợ các truy vấn phức tạp, đảm bảo tính toàn vẹn ACID tuyệt đối.")
    add_para(doc, "Cơ sở lý thuyết về Row Level Security (RLS): RLS là một tính năng bảo mật nâng cao của PostgreSQL, cho phép định nghĩa các chính sách để kiểm soát quyền truy cập dữ liệu ở cấp độ từng hàng (row-level) thay vì cấp độ toàn bộ bảng (table-level). Dựa trên thuộc tính của người dùng đang thực thi câu truy vấn, RLS tự động chèn thêm các điều kiện lọc ẩn vào mệnh đề WHERE của mọi lệnh SQL. Điều này bảo đảm rằng người dùng chỉ có thể nhìn thấy và thao tác trên những dữ liệu mà họ được cấp phép, mang lại một lớp phòng thủ dữ liệu sâu.")

    add_heading2(doc, "2.3.2 Về phía Frontend")

    add_heading3(doc, "2.3.2.1 Thư viện ReactJS (Phiên bản 18)")
    add_para(doc, "Định nghĩa: ReactJS là một thư viện JavaScript mã nguồn mở được phát triển bởi Facebook, chuyên dụng để xây dựng giao diện người dùng có tính tương tác cao.")
    add_para(doc, "Đặc điểm nổi bật: Cơ chế Virtual DOM là một bản sao ảo của cây DOM trên trình duyệt. Khi có sự thay đổi dữ liệu, React sẽ tính toán sự khác biệt giữa Virtual DOM mới và cũ, từ đó chỉ cập nhật chính xác các phần tử thực sự thay đổi trên giao diện thực.")
    add_para(doc, "Ưu điểm: Hiệu năng render xuất sắc, kiến trúc Component dễ tái sử dụng, hệ sinh thái cộng đồng khổng lồ.")
    add_para(doc, "Nhược điểm: React chỉ là thư viện về giao diện, không phải là một framework toàn diện, do đó cần phải cài đặt thêm nhiều công cụ phụ trợ (định tuyến, quản lý trạng thái).")

    add_heading3(doc, "2.3.2.2 Ngôn ngữ TypeScript")
    add_para(doc, "Định nghĩa: TypeScript là một tập hợp siêu ngữ (superset) của JavaScript do Microsoft phát triển, cung cấp thêm hệ thống kiểu dữ liệu tĩnh (static typing) cho JavaScript.")
    add_para(doc, "Đặc điểm nổi bật: Tính năng an toàn kiểu dữ liệu (Type Safety). Lập trình viên phải định nghĩa rõ ràng kiểu dữ liệu cho các biến, hàm và đối tượng. Trình biên dịch TypeScript sẽ thực hiện kiểm tra kiểu trước khi biên dịch mã sang JavaScript thuần.")
    add_para(doc, "Ưu điểm: Giúp phát hiện các lỗi cú pháp và lỗi kiểu dữ liệu ngay trong quá trình viết code (compile-time) thay vì chờ đến khi chạy ứng dụng (runtime), tăng cường sự an toàn và hỗ trợ nhắc code thông minh.")
    add_para(doc, "Nhược điểm: Đòi hỏi lập trình viên phải học thêm cú pháp mới và khai báo kiểu dữ liệu, kéo dài thời gian viết mã ban đầu.")

    add_heading3(doc, "2.3.2.3 Framework Tailwind CSS")
    add_para(doc, "Định nghĩa: Tailwind CSS là một framework CSS theo triết lý Utility-first. Thay vì viết các lớp CSS mang tính ngữ nghĩa trừu tượng, Tailwind cung cấp hàng nghìn lớp tiện ích nhỏ, chức năng đơn lẻ để thiết kế trực tiếp trong mã HTML.")
    add_para(doc, "Đặc điểm nổi bật: Cho phép xây dựng các xu hướng giao diện hiện đại một cách nhanh chóng. Điển hình như Dark Mode (cho phép chuyển đổi màu sắc thành xám/đen thông qua các lớp tiền tố) và Glassmorphism (hiệu ứng kính mờ thông qua việc sử dụng các lớp backdrop-blur và nền bán trong suốt).")
    add_para(doc, "Ưu điểm: Tăng tốc độ thiết kế, không cần phải suy nghĩ đặt tên lớp (class naming), tối ưu dung lượng file CSS ở chế độ production bằng cách loại bỏ các lớp không sử dụng.")
    add_para(doc, "Nhược điểm: Mã HTML có thể trở nên rối mắt và cồng kềnh do chứa quá nhiều lớp CSS.")

    doc.save(filepath)

if __name__ == "__main__":
    generate_document("Chuong2_KienThucNen.docx")
    print("Done")
