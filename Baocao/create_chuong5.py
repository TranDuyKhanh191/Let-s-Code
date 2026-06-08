import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_style(run, bold=False, italic=False, size=13):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_style(run, bold=True, size=15)
    else:
        apply_style(run, bold=True, italic=(level==3), size=13)

def add_paragraph(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(36)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    apply_style(run, size=13, italic=italic, bold=bold)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    apply_style(run, size=13)

def create_chuong5():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    add_heading(doc, "CHƯƠNG V. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 0)
    add_paragraph(doc, "Sau quá trình nghiên cứu, thiết kế, tối ưu hóa và hiện thực hóa đề tài “Xây dựng nền tảng website hỗ trợ giáo viên dạy lập trình cho trẻ em”, dự án đã hoàn thành các mục tiêu đề ra ban đầu, đồng thời phát triển thêm nhiều tính năng đột phá. Dưới đây là tổng kết về các kết quả đạt được, những điểm hạn chế còn tồn tại và phương hướng phát triển của hệ thống trong tương lai.")

    add_heading(doc, "5.1. KẾT QUẢ ĐẠT ĐƯỢC", 1)
    
    add_heading(doc, "5.1.1. Kết quả về mặt kỹ thuật và công nghệ", 2)
    add_paragraph(doc, "Hệ thống Cơ sở dữ liệu và Bảo mật (Database & Security): Thiết kế và triển khai thành công lược đồ cơ sở dữ liệu quan hệ trên PostgreSQL (Supabase) đảm bảo tính toàn vẹn 3NF. Dự án đã kích hoạt và thiết lập thành công chính sách bảo mật cấp dữ liệu (Row Level Security - RLS) để phân quyền chặt chẽ giữa các vai trò (Admin, Teacher, Student), ngăn chặn truy cập trái phép. Bên cạnh đó, tính năng Realtime (WebSocket) được áp dụng giúp đồng bộ thông báo cập nhật phân công theo thời gian thực.")
    add_paragraph(doc, "Hệ thống Backend & API: Tối ưu hóa hiệu suất truy xuất dữ liệu bằng kỹ thuật Aggregation (JSONB) thông qua PostgreSQL Functions, giúp giảm đáng kể độ trễ (latency) khi tải nội dung bài học. Cơ chế bảo mật xác thực (Authentication) dựa trên JSON Web Token (JWT) được triển khai ổn định và an toàn.")
    add_paragraph(doc, "Ứng dụng Trí tuệ Nhân tạo và Trải nghiệm Tương tác: Lần đầu tiên tích hợp thành công mô hình Gemini Vision (AI Tutor) vào phân hệ của học sinh, cho phép AI phân tích hình ảnh khối lệnh lỗi và đưa ra gợi ý sư phạm. Công nghệ Server-Sent Events (SSE) được áp dụng giúp trả về luồng phản hồi của AI liên tục như một đoạn chat thực. Bên cạnh đó, hiệu ứng khen thưởng trực quan (Gamification với canvas-confetti) được tích hợp để khích lệ học sinh khi nộp bài thành công.")
    add_paragraph(doc, "Giao diện người dùng (Frontend): Xây dựng trên nền tảng ReactJS, TypeScript và Tailwind CSS, mang lại giao diện Glassmorphism hiện đại, tối ưu hóa Responsive. Trình phát bài giảng (Lesson Player) ứng dụng cơ chế render thông minh không cần tải lại trang, tích hợp đa phương tiện (Video HD, tài liệu PDF) mượt mà.")
    
    add_heading(doc, "5.1.2. Kết quả về mặt nghiệp vụ và quản lý", 2)
    add_paragraph(doc, "Số hóa quy trình quản trị: Hệ thống cung cấp cho Admin và Staff một quy trình quản trị tập trung, thay thế hoàn toàn việc quản lý giáo án bằng các file rời rạc. Dữ liệu chương trình, khóa học, bài học và phân công giảng dạy được quản lý khoa học với cơ chế phê duyệt chặt chẽ.")
    add_paragraph(doc, "Phân hệ Cổng Giáo viên (Teacher Portal): Xóa bỏ khó khăn trong việc chuẩn bị giáo án. Trình phát bài giảng hỗ trợ giáo viên thuyết trình tại lớp một cách trực quan, có đầy đủ công cụ chuyển tiếp nội dung, câu hỏi trắc nghiệm (Quiz). Hệ thống còn cho phép giáo viên theo dõi tiến độ của lớp, nhận và tải về file dự án (.llsp3) của học sinh một cách dễ dàng.")
    add_paragraph(doc, "Phân hệ Cổng Tự học (Student Portal): Tạo ra một không gian học tập riêng biệt, an toàn cho trẻ em. Nội dung bài giảng được tái sử dụng nhưng loại bỏ các quyền quản trị, giúp học sinh có thể ôn lại kiến thức, kéo thả (Drag & Drop) nộp bài trực tuyến, và đặc biệt là nhận sự hỗ trợ 24/7 từ Trợ lý AI khi mắc lỗi lập trình Robot.")

    add_heading(doc, "5.2. HẠN CHẾ CỦA ĐỀ TÀI", 1)
    add_paragraph(doc, "Mặc dù đã hoàn thành các tính năng theo yêu cầu và đề xuất, hệ thống vẫn còn một số giới hạn nhất định trong phạm vi đồ án:")
    add_bullet(doc, "Phạm vi nội dung đào tạo: Hệ thống hiện tại đang tập trung thiết kế chuyên biệt cho chương trình Robotics (SPIKE Essential và SPIKE Prime). Các dạng bài tập đặc thù của các môn học khác (như Scratch, Python) chưa được tích hợp hệ thống chấm điểm tự động.")
    add_bullet(doc, "Giới hạn của Mô hình AI: Trợ lý AI (Vision Tutor) hoạt động phụ thuộc vào chất lượng ảnh chụp màn hình khối lệnh của học sinh. Nếu ảnh bị mờ, cắt xén, hoặc quá nhiều khối lệnh phức tạp, độ chính xác trong phân tích của Gemini có thể giảm. Tốc độ luồng phản hồi SSE cũng phụ thuộc vào kết nối API bên ngoài.")
    add_bullet(doc, "Khả năng chịu tải đồng thời: Các tính năng dùng WebSocket (Realtime) và SSE (Stream AI) mới được kiểm thử trong môi trường có lưu lượng vừa phải. Cần có thêm các kịch bản kiểm thử (Stress Test) khi có hàng chục lớp học, hàng trăm học sinh sử dụng đồng thời.")
    add_bullet(doc, "Nền tảng ứng dụng: Hiện tại sản phẩm hoàn toàn là Web Application. Mặc dù giao diện đã tương thích di động (Responsive), nhưng chưa có phiên bản Native Mobile App độc lập (iOS/Android) để mang lại trải nghiệm truy cập tốt nhất trên tablet cho trẻ em.")

    add_heading(doc, "5.3. HƯỚNG PHÁT TRIỂN TRONG TƯƠNG LAI", 1)
    add_paragraph(doc, "Dựa trên kết quả đạt được, hệ thống có nhiều tiềm năng mở rộng để trở thành một nền tảng EdTech toàn diện:")
    add_bullet(doc, "Hệ thống Học tập Thích ứng (Adaptive Learning): Tận dụng mô hình AI (Agentic AI) để phân tích lịch sử làm bài (Quiz) và các lỗi thường gặp của học sinh, từ đó tự động gợi ý các khóa học bổ trợ hoặc điều chỉnh độ khó của thử thách (Challenge) phù hợp với năng lực từng cá nhân.")
    add_bullet(doc, "Tích hợp đa nền tảng và Chấm điểm tự động (Auto-grading): Mở rộng hỗ trợ nộp bài và chạy thử mã nguồn (Sandbox) trực tiếp trên trình duyệt đối với các ngôn ngữ khác như Python, JavaScript, thay vì chỉ upload file dự án (.llsp3).")
    add_bullet(doc, "Phát triển hệ thống Gamification toàn diện: Thay vì chỉ có hiệu ứng khen thưởng cục bộ, hệ thống sẽ xây dựng hệ thống điểm thưởng (Points), cấp độ (Levels), và bảng xếp hạng (Leaderboards/Badges) nhằm thúc đẩy sự thi đua học tập giữa các học sinh trong trung tâm.")
    add_bullet(doc, "Tương tác thời gian thực tại lớp (Classroom Interaction): Phát triển các tính năng như 'Giơ tay phát biểu', 'Làm bài Quiz đồng bộ' (giống Kahoot) hoặc nhắn tin trực tiếp với giáo viên ngay trên Student Portal, biến web app thành một công cụ điều phối lớp học thực thụ.")
    add_bullet(doc, "Ứng dụng Di động (Mobile App) cho Phụ huynh: Xây dựng cổng thông tin trên ứng dụng di động dành riêng cho phụ huynh để theo dõi tiến độ học, điểm số, xem các dự án Robot đã hoàn thành của con em mình và nhận thông báo từ trung tâm.")

    output_path = os.path.join(os.path.dirname(__file__), 'Chuong5_KetLuan.docx')
    doc.save(output_path)
    print(f"File created successfully at {output_path}")

    # Generate Markdown file as well
    md_output_path = os.path.join(os.path.dirname(__file__), 'Chuong5_KetLuan.md')
    with open(md_output_path, 'w', encoding='utf-8') as f:
        f.write("# CHƯƠNG V. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN\n\n")
        f.write("Sau quá trình nghiên cứu, thiết kế, tối ưu hóa và hiện thực hóa đề tài “Xây dựng nền tảng website hỗ trợ giáo viên dạy lập trình cho trẻ em”, dự án đã hoàn thành các mục tiêu đề ra ban đầu, đồng thời phát triển thêm nhiều tính năng đột phá. Dưới đây là tổng kết về các kết quả đạt được, những điểm hạn chế còn tồn tại và phương hướng phát triển của hệ thống trong tương lai.\n\n")
        
        f.write("## 5.1. KẾT QUẢ ĐẠT ĐƯỢC\n\n")
        
        f.write("### 5.1.1. Kết quả về mặt kỹ thuật và công nghệ\n")
        f.write("- **Hệ thống Cơ sở dữ liệu và Bảo mật (Database & Security):** Thiết kế và triển khai thành công lược đồ cơ sở dữ liệu quan hệ trên PostgreSQL (Supabase) đảm bảo tính toàn vẹn 3NF. Dự án đã kích hoạt và thiết lập thành công chính sách bảo mật cấp dữ liệu (Row Level Security - RLS) để phân quyền chặt chẽ giữa các vai trò (Admin, Teacher, Student), ngăn chặn truy cập trái phép. Bên cạnh đó, tính năng Realtime (WebSocket) được áp dụng giúp đồng bộ thông báo cập nhật phân công theo thời gian thực.\n")
        f.write("- **Hệ thống Backend & API:** Tối ưu hóa hiệu suất truy xuất dữ liệu bằng kỹ thuật Aggregation (JSONB) thông qua PostgreSQL Functions, giúp giảm đáng kể độ trễ (latency) khi tải nội dung bài học. Cơ chế bảo mật xác thực (Authentication) dựa trên JSON Web Token (JWT) được triển khai ổn định và an toàn.\n")
        f.write("- **Ứng dụng Trí tuệ Nhân tạo và Trải nghiệm Tương tác:** Lần đầu tiên tích hợp thành công mô hình Gemini Vision (AI Tutor) vào phân hệ của học sinh, cho phép AI phân tích hình ảnh khối lệnh lỗi và đưa ra gợi ý sư phạm. Công nghệ Server-Sent Events (SSE) được áp dụng giúp trả về luồng phản hồi của AI liên tục như một đoạn chat thực. Bên cạnh đó, hiệu ứng khen thưởng trực quan (Gamification với canvas-confetti) được tích hợp để khích lệ học sinh khi nộp bài thành công.\n")
        f.write("- **Giao diện người dùng (Frontend):** Xây dựng trên nền tảng ReactJS, TypeScript và Tailwind CSS, mang lại giao diện Glassmorphism hiện đại, tối ưu hóa Responsive. Trình phát bài giảng (Lesson Player) ứng dụng cơ chế render thông minh không cần tải lại trang, tích hợp đa phương tiện (Video HD, tài liệu PDF) mượt mà.\n\n")
        
        f.write("### 5.1.2. Kết quả về mặt nghiệp vụ và quản lý\n")
        f.write("- **Số hóa quy trình quản trị:** Hệ thống cung cấp cho Admin và Staff một quy trình quản trị tập trung, thay thế hoàn toàn việc quản lý giáo án bằng các file rời rạc. Dữ liệu chương trình, khóa học, bài học và phân công giảng dạy được quản lý khoa học với cơ chế phê duyệt chặt chẽ.\n")
        f.write("- **Phân hệ Cổng Giáo viên (Teacher Portal):** Xóa bỏ khó khăn trong việc chuẩn bị giáo án. Trình phát bài giảng hỗ trợ giáo viên thuyết trình tại lớp một cách trực quan, có đầy đủ công cụ chuyển tiếp nội dung, câu hỏi trắc nghiệm (Quiz). Hệ thống còn cho phép giáo viên theo dõi tiến độ của lớp, nhận và tải về file dự án (.llsp3) của học sinh một cách dễ dàng.\n")
        f.write("- **Phân hệ Cổng Tự học (Student Portal):** Tạo ra một không gian học tập riêng biệt, an toàn cho trẻ em. Nội dung bài giảng được tái sử dụng nhưng loại bỏ các quyền quản trị, giúp học sinh có thể ôn lại kiến thức, kéo thả (Drag & Drop) nộp bài trực tuyến, và đặc biệt là nhận sự hỗ trợ 24/7 từ Trợ lý AI khi mắc lỗi lập trình Robot.\n\n")
        
        f.write("## 5.2. HẠN CHẾ CỦA ĐỀ TÀI\n\n")
        f.write("Mặc dù đã hoàn thành các tính năng theo yêu cầu và đề xuất, hệ thống vẫn còn một số giới hạn nhất định trong phạm vi đồ án:\n")
        f.write("- **Phạm vi nội dung đào tạo:** Hệ thống hiện tại đang tập trung thiết kế chuyên biệt cho chương trình Robotics (SPIKE Essential và SPIKE Prime). Các dạng bài tập đặc thù của các môn học khác (như Scratch, Python) chưa được tích hợp hệ thống chấm điểm tự động.\n")
        f.write("- **Giới hạn của Mô hình AI:** Trợ lý AI (Vision Tutor) hoạt động phụ thuộc vào chất lượng ảnh chụp màn hình khối lệnh của học sinh. Nếu ảnh bị mờ, cắt xén, hoặc quá nhiều khối lệnh phức tạp, độ chính xác trong phân tích của Gemini có thể giảm. Tốc độ luồng phản hồi SSE cũng phụ thuộc vào kết nối API bên ngoài.\n")
        f.write("- **Khả năng chịu tải đồng thời:** Các tính năng dùng WebSocket (Realtime) và SSE (Stream AI) mới được kiểm thử trong môi trường có lưu lượng vừa phải. Cần có thêm các kịch bản kiểm thử (Stress Test) khi có hàng chục lớp học, hàng trăm học sinh sử dụng đồng thời.\n")
        f.write("- **Nền tảng ứng dụng:** Hiện tại sản phẩm hoàn toàn là Web Application. Mặc dù giao diện đã tương thích di động (Responsive), nhưng chưa có phiên bản Native Mobile App độc lập (iOS/Android) để mang lại trải nghiệm truy cập tốt nhất trên tablet cho trẻ em.\n\n")
        
        f.write("## 5.3. HƯỚNG PHÁT TRIỂN TRONG TƯƠNG LAI\n\n")
        f.write("Dựa trên kết quả đạt được, hệ thống có nhiều tiềm năng mở rộng để trở thành một nền tảng EdTech toàn diện:\n")
        f.write("- **Hệ thống Học tập Thích ứng (Adaptive Learning):** Tận dụng mô hình AI (Agentic AI) để phân tích lịch sử làm bài (Quiz) và các lỗi thường gặp của học sinh, từ đó tự động gợi ý các khóa học bổ trợ hoặc điều chỉnh độ khó của thử thách (Challenge) phù hợp với năng lực từng cá nhân.\n")
        f.write("- **Tích hợp đa nền tảng và Chấm điểm tự động (Auto-grading):** Mở rộng hỗ trợ nộp bài và chạy thử mã nguồn (Sandbox) trực tiếp trên trình duyệt đối với các ngôn ngữ khác như Python, JavaScript, thay vì chỉ upload file dự án (.llsp3).\n")
        f.write("- **Phát triển hệ thống Gamification toàn diện:** Thay vì chỉ có hiệu ứng khen thưởng cục bộ, hệ thống sẽ xây dựng hệ thống điểm thưởng (Points), cấp độ (Levels), và bảng xếp hạng (Leaderboards/Badges) nhằm thúc đẩy sự thi đua học tập giữa các học sinh trong trung tâm.\n")
        f.write("- **Tương tác thời gian thực tại lớp (Classroom Interaction):** Phát triển các tính năng như 'Giơ tay phát biểu', 'Làm bài Quiz đồng bộ' (giống Kahoot) hoặc nhắn tin trực tiếp với giáo viên ngay trên Student Portal, biến web app thành một công cụ điều phối lớp học thực thụ.\n")
        f.write("- **Ứng dụng Di động (Mobile App) cho Phụ huynh:** Xây dựng cổng thông tin trên ứng dụng di động dành riêng cho phụ huynh để theo dõi tiến độ học, điểm số, xem các dự án Robot đã hoàn thành của con em mình và nhận thông báo từ trung tâm.\n")

    print(f"Markdown file created successfully at {md_output_path}")

if __name__ == '__main__':
    create_chuong5()
