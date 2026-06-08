import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_style(run, bold=False, italic=False, size=13):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    if level == 0: 
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_style(run, bold=True, size=15)
    elif level == 1: 
        apply_style(run, bold=True, size=14)
    elif level == 2: 
        apply_style(run, bold=True, size=13)

def add_paragraph(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    apply_style(run, size=13, italic=italic, bold=bold)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    apply_style(run, size=13)

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(36)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(11)

def create_doc():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    add_heading(doc, "HƯỚNG DẪN VẼ SƠ ĐỒ VÀ CHÈN ẢNH CODE VÀO BÁO CÁO", 0)

    add_paragraph(doc, "Để bài báo cáo trực quan và sinh động hơn, bạn hãy thực hiện theo các bước dưới đây để vẽ sơ đồ và chụp ảnh màn hình code chèn vào file Word.")

    add_heading(doc, "PHẦN 1: HƯỚNG DẪN VẼ SƠ ĐỒ BẰNG MERMAID", 1)
    add_paragraph(doc, "Bạn hãy truy cập vào trang web: https://mermaid.live/")
    add_paragraph(doc, "Sau đó copy các đoạn code dưới đây, dán vào cột \"Code\" bên trái của trang web. Trang web sẽ tự động vẽ ra sơ đồ bên phải. Bạn hãy lưu ảnh (hoặc chụp màn hình sơ đồ) và dán vào file Word.")

    add_heading(doc, "1. Sơ đồ ERD (Entity Relationship Diagram) Đầy đủ", 2)
    add_bullet(doc, "Vị trí chèn trong Word: Dưới mục 3.2.2 Tổng quan cấu trúc cơ sở dữ liệu")
    add_bullet(doc, "Caption hình ảnh: Hình 3.1: Sơ đồ mối quan hệ thực thể (ERD) chi tiết các bảng của hệ thống")
    add_paragraph(doc, "Code Mermaid (Copy nội dung bên dưới):")
    erd_code = """erDiagram
    users {
        bigint id PK
        text email
        text password
        text role
        text full_name
        timestamp created_at
    }
    programs {
        bigint id PK
        text name
        text status
    }
    courses {
        bigint id PK
        bigint program_id FK
        text name
        text course_code
    }
    lessons {
        bigint id PK
        bigint course_id FK
        text title
        text status
    }
    assignments {
        bigint id PK
        bigint teacher_id FK
        text resource_type
        int resource_id
    }
    enrollments {
        bigint id PK
        bigint student_id FK
        bigint course_id FK
        text status
    }
    lesson_progress {
        bigint id PK
        bigint student_id FK
        bigint lesson_id FK
        text status
        text submitted_file_url
        int score
    }
    lesson_contents {
        bigint id PK
        bigint lesson_id FK
        text title
        text description
    }
    lesson_objectives {
        bigint id PK
        bigint lesson_id FK
        text knowledge
        text skills
    }
    lesson_quizzes {
        bigint id PK
        bigint lesson_id FK
        text question_text
        text quiz_type
    }
    quiz_answers {
        bigint id PK
        bigint quiz_id FK
        text answer_text
        boolean is_correct
    }
    media {
        bigint id PK
        text url
        text mime_type
    }
    content_media {
        bigint id PK
        bigint media_id FK
        text content_type
        bigint content_id
    }

    users ||--o{ enrollments : "đăng ký"
    users ||--o{ assignments : "được phân công"
    users ||--o{ lesson_progress : "theo dõi tiến độ"
    programs ||--o{ courses : "bao gồm"
    courses ||--o{ lessons : "bao gồm"
    courses ||--o{ enrollments : "có"
    lessons ||--o{ lesson_contents : "chứa"
    lessons ||--o{ lesson_objectives : "chứa"
    lessons ||--o{ lesson_quizzes : "chứa"
    lessons ||--o{ lesson_progress : "được theo dõi"
    lesson_quizzes ||--o{ quiz_answers : "có"
    media ||--o{ content_media : "đính kèm"
"""
    add_code_block(doc, erd_code)

    add_heading(doc, "2. Sơ đồ DFD (Data Flow Diagram) Mức 1 - Phân rã chức năng", 2)
    add_bullet(doc, "Vị trí chèn trong Word: Dưới mục 3.1.1 Mục tiêu của hệ thống web bài giảng")
    add_bullet(doc, "Caption hình ảnh: Hình 3.2: Sơ đồ luồng dữ liệu (DFD) mức 1 - Chi tiết các tiến trình và kho dữ liệu")
    add_paragraph(doc, "Code Mermaid (Copy nội dung bên dưới):")
    dfd_code = """flowchart TD
    GV[Giáo viên]
    HS[Học sinh]
    Admin[Quản trị viên]
    
    P1((1. Quản lý\\nTài khoản))
    P2((2. Quản lý\\nKhóa học))
    P3((3. Học tập &\\nTương tác))
    P4((4. Đánh giá\\n& Báo cáo))
    
    D1[(D1: Bảng Users)]
    D2[(D2: Bảng Courses)]
    D3[(D3: Bảng Progress)]
    D4[(D4: Bảng Assignments)]
    
    Admin -->|Thông tin TK| P1
    P1 <-->|Lưu/Lấy User| D1
    
    Admin -->|Tạo Khóa học| P2
    P2 <-->|Lưu/Lấy Course| D2
    Admin -->|Phân công GV| P2
    P2 <-->|Lưu Phân công| D4
    
    GV -->|Quản lý Nội dung| P2
    GV -->|Chấm điểm, Feedback| P4
    
    HS -->|Truy cập bài học| P3
    P3 <-->|Truy xuất Course| D2
    HS -->|Nộp bài| P3
    
    P3 -->|Cập nhật tiến độ| P4
    P4 <-->|Lưu/Lấy Tiến độ| D3
"""
    add_code_block(doc, dfd_code)

    add_heading(doc, "3. Sơ đồ PFD (Process Flow Diagram) - Chi tiết luồng học tập", 2)
    add_bullet(doc, "Vị trí chèn trong Word: Dưới mục 3.7.7 Trình phát bài giảng tương tác (Interactive Lesson Player)")
    add_bullet(doc, "Caption hình ảnh: Hình 3.3: Lưu đồ thuật toán quy trình học tập chi tiết của học sinh")
    add_paragraph(doc, "Code Mermaid (Copy nội dung bên dưới):")
    pfd_code = """flowchart TD
    A([Bắt đầu]) --> B[Đăng nhập]
    B --> C{Xác thực user?}
    C -->|Sai| B
    C -->|Đúng| D[Xem Dashboard]
    D --> E[Chọn Khóa học & Bài học]
    E --> F[Xem Video / Tài liệu]
    F --> G{Có Bài tập/\\nThử thách?}
    G -->|Không| K[Đánh dấu hoàn thành]
    G -->|Có| I[Thực hành / Trắc nghiệm]
    I --> J{Cần trợ giúp?}
    J -->|Có| AI[Chat với AI Tutor]
    AI --> I
    J -->|Không| NOP[Nộp bài]
    NOP --> chấm{Hệ thống / GV\\nChấm điểm}
    chấm -->|Chưa đạt| I
    chấm -->|Đạt| K
    K --> L([Kết thúc bài học])
"""
    add_code_block(doc, pfd_code)

    add_heading(doc, "PHẦN 2: DANH SÁCH CÁC ĐOẠN CODE CẦN CHỤP ẢNH VÀ CHÈN", 1)
    add_paragraph(doc, "Bạn hãy mở phần mềm VS Code, mở các file dưới đây, khoanh vùng đoạn code quan trọng, chụp màn hình và dán vào Word theo đúng vị trí hướng dẫn.")

    add_heading(doc, "Ảnh Code 1: Khởi tạo kiến trúc Store / Context", 2)
    add_bullet(doc, "Vị trí chèn: Dưới mục 3.4.1 Mô hình tổ chức mã nguồn")
    add_bullet(doc, "Mở file: frontend/src/context/AuthContext.tsx")
    add_bullet(doc, "Vùng chụp: Phần định nghĩa AuthContext và useAuth (Tầm 15 - 20 dòng đầu).")
    add_bullet(doc, "Caption (Ghi chú): Hình 3.4: Khởi tạo Context API để quản lý trạng thái xác thực người dùng")

    add_heading(doc, "Ảnh Code 2: Xử lý gộp dữ liệu khóa học (Data Aggregation)", 2)
    add_bullet(doc, "Vị trí chèn: Dưới đoạn chữ '...thay vì gọi API tuần tự, sử dụng Promise.all để kích hoạt các request song song...' ở mục 3.7.1.1")
    add_bullet(doc, "Mở file: frontend/src/pages/admin/courses/AdminCoursesPage.tsx")
    add_bullet(doc, "Vùng chụp: Bên trong hàm fetchCourses(), đoạn gọi Promise.all([api.get('/programs/1/courses'), ...]).")
    add_bullet(doc, "Caption (Ghi chú): Hình 3.5: Thuật toán sử dụng Promise.all để gộp dữ liệu khóa học từ nhiều chương trình")

    add_heading(doc, "Ảnh Code 3: Component AssignmentModal (Phòng tránh trùng lặp)", 2)
    add_bullet(doc, "Vị trí chèn: Dưới mục 3.7.3 Phân hệ quản lý nhân sự và phân công")
    add_bullet(doc, "Mở file: frontend/src/pages/admin/teachers/components/AssignmentModal.tsx")
    add_bullet(doc, "Vùng chụp: Đoạn logic sử dụng hàm filter và some để loại bỏ các khóa học đã được phân công.")
    add_bullet(doc, "Caption (Ghi chú): Hình 3.6: Logic xử lý mảng để ngăn chặn phân công trùng lặp khóa học")

    add_heading(doc, "Ảnh Code 4: Cập nhật giao diện lạc quan (Optimistic UI Update)", 2)
    add_bullet(doc, "Vị trí chèn: Dưới mục 3.6.3 Quản lý trạng thái và luồng dữ liệu")
    add_bullet(doc, "Mở file: frontend/src/pages/admin/teachers/ManageTeachersPage.tsx")
    add_bullet(doc, "Vùng chụp: Trong hàm handleDeleteTeacher(), đoạn mã setTeachers(teachers.filter(t => t.id !== id)) ngay trước hoặc trong khi gọi API.")
    add_bullet(doc, "Caption (Ghi chú): Hình 3.7: Cập nhật giao diện tức thì (Optimistic UI) khi xóa bản ghi")

    add_heading(doc, "Ảnh Code 5: Xử lý Backend Service (Tầng logic)", 2)
    add_bullet(doc, "Vị trí chèn: Dưới mục 3.3.4 Mô tả chi tiết các service backend đã xây dựng")
    add_bullet(doc, "Mở file: backend/src/services/course.service.ts")
    add_bullet(doc, "Vùng chụp: Bất kỳ một hàm lớn nào (ví dụ: createCourse hoặc updateCourse) có xử lý logic Supabase.")
    add_bullet(doc, "Caption (Ghi chú): Hình 3.8: Cấu trúc xử lý nghiệp vụ tại tầng Service của Backend")

    add_heading(doc, "Ảnh Code 6: Chức năng Chatbot AI (AI Tutor)", 2)
    add_bullet(doc, "Vị trí chèn: Dưới mục 3.7.7 Trình phát bài giảng tương tác")
    add_bullet(doc, "Mở file: frontend/src/pages/student/ChallengeSandbox.tsx")
    add_bullet(doc, "Vùng chụp: Đoạn mã xử lý gọi API gửi tin nhắn tới Gemini (/api/ai/chat).")
    add_bullet(doc, "Caption (Ghi chú): Hình 3.9: Logic gọi API gửi tin nhắn và nhận phản hồi từ AI Tutor")

    output_path = "GhiChu_ChenAnh_V2.docx"
    doc.save(output_path)
    print(f"File created successfully at {output_path}")

if __name__ == "__main__":
    create_doc()
